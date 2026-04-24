"""Parameterized fixup battery — generalized from moore_build/src/fixup_moore.py.

Applies the set of post-paste corrections that Word's "Keep Source Formatting"
makes necessary:
  - style definition updates (body ls, captions Georgia, footnote ls, block
    quote, list spacing, heading keep-next)
  - block-quote remapping (indented body paragraphs → TME Block Quote)
  - caption reclassification (leading-word heuristic)
  - list direct-spacing clear
  - table centering + cantSplit + tblHeader
  - masthead grid rewrite (defensive, in case Word merged adjacent tables)
  - reference run format strip
  - footnote font + size normalization (zip-level edit of footnotes.xml)
"""
from __future__ import annotations

import os
import re
import shutil
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from tme_template.colors import BLOCKQUOTE_INK


# ---------- style definition updates ----------

def _find_style(styles, *candidates):
    names_lc = {s.name.lower(): s.name for s in styles}
    for c in candidates:
        actual = names_lc.get(c.lower())
        if actual is not None:
            return styles[actual]
    return None


def update_styles(doc) -> None:
    styles = doc.styles
    style_names = {s.name for s in styles}

    body = styles["TME Body"]
    body.paragraph_format.line_spacing = 1.25

    for hname in ("TME H1", "TME H2", "TME H3"):
        if hname in style_names:
            h = styles[hname]
            h.paragraph_format.keep_with_next = True
            h.paragraph_format.keep_together = True

    fc = styles["TME Figure Caption"]
    fc.font.name = "Georgia"
    fc.paragraph_format.keep_with_next = True  # caption above → glue down to figure/table

    tc = styles["TME Table Caption"]
    tc.font.name = "Georgia"
    tc.paragraph_format.keep_with_next = True

    fn = styles["TME Footnote"]
    fn.font.name = "Georgia"
    fn.paragraph_format.line_spacing = 1.2

    # Word's built-in footnote styles that Keep Source Formatting imports
    ft = _find_style(styles, "Footnote Text", "footnote text")
    if ft is not None:
        ft.font.name = "Georgia"
        ft.font.size = Pt(9)
        if ft.paragraph_format is not None:
            ft.paragraph_format.line_spacing = 1.2
    for sname in ("Footnote Reference", "footnote reference", "Footnote Text Char"):
        s = _find_style(styles, sname)
        if s is not None:
            s.font.name = "Georgia"

    if "TME Block Quote" in style_names:
        bq = styles["TME Block Quote"]
    else:
        bq = styles.add_style("TME Block Quote", WD_STYLE_TYPE.PARAGRAPH)
    bq.font.name = "Georgia"
    bq.font.size = Pt(10.5)
    bq.font.color.rgb = RGBColor.from_string(BLOCKQUOTE_INK)
    bq.paragraph_format.left_indent = Pt(28)
    bq.paragraph_format.right_indent = Pt(28)
    bq.paragraph_format.line_spacing = 1.0
    bq.paragraph_format.space_before = Pt(8)
    bq.paragraph_format.space_after = Pt(8)

    if "List Paragraph" in style_names:
        lp = styles["List Paragraph"]
        lp.font.name = "Georgia"
        lp.font.size = Pt(11.5)
        lp.paragraph_format.line_spacing = 1.25
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after = Pt(4)


# ---------- content remapping ----------

_REF_OPENER = re.compile(r"^[A-ZÀ-ÖØ-Ý][\w'’\-]+,\s+[A-Z]\.")


def remap_block_quotes(doc) -> int:
    bq_style = doc.styles["TME Block Quote"]
    n = 0
    for p in doc.paragraphs:
        if p.style.name != "TME Body":
            continue
        # Guard: a "LastName, F." opener means this is a reference that was
        # mis-classified as body; don't turn it into a block quote.
        if _REF_OPENER.match(p.text.strip()):
            continue
        # Guard: hanging indent (negative first_line_indent) strongly suggests
        # a reference, not a block quote
        fli = p.paragraph_format.first_line_indent
        if fli is not None and fli.pt < 0:
            continue
        li = p.paragraph_format.left_indent
        if li is None:
            continue
        if li.pt > 20:
            p.style = bq_style
            p.paragraph_format.left_indent = None
            p.paragraph_format.first_line_indent = None
            n += 1
    return n


def rescue_misclassified_references(doc) -> int:
    """Move any TME Block Quote whose text looks like a reference opener back
    to TME Reference. Defensive against upstream misclassification."""
    ref_style = doc.styles["TME Reference"]
    n = 0
    for p in doc.paragraphs:
        if p.style.name != "TME Block Quote":
            continue
        if _REF_OPENER.match(p.text.strip()):
            p.style = ref_style
            p.paragraph_format.left_indent = None
            p.paragraph_format.first_line_indent = None
            n += 1
    return n


def fix_caption_classifications(doc) -> dict:
    fc = doc.styles["TME Figure Caption"]
    tc = doc.styles["TME Table Caption"]
    stats = {"fig_from_body": 0, "tab_from_body": 0, "fig_fix": 0, "tab_fix": 0}
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t or len(t) > 400:
            continue
        starts_fig = t.startswith("Figure")
        starts_tab = t.startswith("Table")
        if not (starts_fig or starts_tab):
            continue
        sn = p.style.name
        if starts_fig and sn == "TME Body":
            p.style = fc
            stats["fig_from_body"] += 1
        elif starts_tab and sn == "TME Body":
            p.style = tc
            stats["tab_from_body"] += 1
        elif starts_fig and sn == "TME Table Caption":
            p.style = fc
            stats["fig_fix"] += 1
        elif starts_tab and sn == "TME Figure Caption":
            p.style = tc
            stats["tab_fix"] += 1
    return stats


def clear_list_direct_spacing(doc) -> int:
    n = 0
    for p in doc.paragraphs:
        if p.style.name != "List Paragraph":
            continue
        pf = p.paragraph_format
        changed = False
        if pf.line_spacing is not None:
            pf.line_spacing = None
            changed = True
        if pf.space_before is not None and pf.space_before.pt > 0:
            pf.space_before = None
            changed = True
        if changed:
            n += 1
    return n


def strip_reference_run_formatting(doc) -> int:
    n = 0
    for p in doc.paragraphs:
        if p.style.name != "TME Reference":
            continue
        for r in p.runs:
            rPr = r._r.find(qn("w:rPr"))
            if rPr is None:
                continue
            for tag in ("w:sz", "w:szCs", "w:b", "w:bCs", "w:rFonts"):
                for el in rPr.findall(qn(tag)):
                    rPr.remove(el)
                    n += 1
    return n


# Paragraph-property tags to strip from TME-styled paragraphs. These are the
# direct overrides that leak through "Keep Source Formatting" paste and fight
# our paragraph styles.
_PPR_STRIP_TAGS = ("w:spacing", "w:ind")

# Run-property tags to strip on structural styles (headings, captions, body,
# references). Leave w:i (italic), w:iCs, w:u (underline), w:color alone —
# those are legitimate inline emphasis.
_RPR_STRIP_TAGS_STRUCT = ("w:rFonts", "w:sz", "w:szCs", "w:b", "w:bCs")


def strip_direct_formatting(doc) -> dict:
    """Strip direct paragraph-level spacing/indent and run-level font/size/bold
    on paragraphs in TME styles where the style should win. Applied to:
      - TME Body, TME H1/H2/H3, TME Figure Caption, TME Table Caption,
        TME Reference, TME Footnote, TME Block Quote, List Paragraph.
    """
    targets = {
        "TME Body", "TME H1", "TME H2", "TME H3",
        "TME Figure Caption", "TME Table Caption",
        "TME Reference", "TME Footnote", "TME Block Quote",
        "List Paragraph",
    }
    stats = {"paras_stripped": 0, "runs_stripped": 0}
    for p in doc.paragraphs:
        if p.style.name not in targets:
            continue
        # Paragraph-level strip
        pPr = p._p.find(qn("w:pPr"))
        para_changed = False
        if pPr is not None:
            for tag in _PPR_STRIP_TAGS:
                for el in pPr.findall(qn(tag)):
                    pPr.remove(el)
                    para_changed = True
        # Run-level strip
        run_changed = False
        for r in p.runs:
            rPr = r._r.find(qn("w:rPr"))
            if rPr is None:
                continue
            for tag in _RPR_STRIP_TAGS_STRUCT:
                for el in rPr.findall(qn(tag)):
                    rPr.remove(el)
                    run_changed = True
        if para_changed:
            stats["paras_stripped"] += 1
        if run_changed:
            stats["runs_stripped"] += 1
    return stats


def report_below_element_captions(doc) -> list:
    """Return caption paragraphs that sit below their figure/table instead of above.

    APA 7 places the caption above the element; fixup flags any violation so
    the editor can fix it (manually, or via swap_captions_above).

    Returned entries look like:
        {"index": 12, "kind": "figure", "preview": "Figure 2. Student work…"}
    """
    reports = []
    body = doc.element.body
    children = list(body.iterchildren())
    tag_p, tag_tbl = qn("w:p"), qn("w:tbl")

    para_idx = -1  # running index into doc.paragraphs (paragraph children only)
    for i, el in enumerate(children):
        if el.tag != tag_p:
            continue
        para_idx += 1
        p = doc.paragraphs[para_idx]
        sn = p.style.name if p.style is not None else ""
        if sn not in ("TME Figure Caption", "TME Table Caption"):
            continue
        # Look at the previous non-empty sibling in body order.
        prev_el = None
        for back in range(i - 1, -1, -1):
            cand = children[back]
            if cand.tag == tag_p:
                # Skip empty paragraph spacers — but a paragraph carrying a
                # drawing/pict is NOT empty for our purposes, even if it has
                # no w:t text.
                has_drawing = cand.find(".//" + qn("w:drawing")) is not None
                has_pict = cand.find(".//" + qn("w:pict")) is not None
                if (
                    not has_drawing
                    and not has_pict
                    and (cand.text or "").strip() == ""
                    and cand.find(".//" + qn("w:t")) is None
                ):
                    continue
                prev_el = cand
                break
            if cand.tag == tag_tbl:
                prev_el = cand
                break
        if prev_el is None:
            continue
        is_fig_caption = sn == "TME Figure Caption"
        is_tab_caption = sn == "TME Table Caption"
        prev_is_image_para = prev_el.tag == tag_p and (
            prev_el.find(".//" + qn("w:drawing")) is not None or
            prev_el.find(".//" + qn("w:pict")) is not None
        )
        prev_is_table = prev_el.tag == tag_tbl
        if is_fig_caption and prev_is_image_para:
            reports.append({
                "index": para_idx, "kind": "figure",
                "preview": (p.text[:80] or "").strip(),
            })
        elif is_tab_caption and prev_is_table:
            reports.append({
                "index": para_idx, "kind": "table",
                "preview": (p.text[:80] or "").strip(),
            })
    return reports


def swap_captions_above(doc, report: list) -> int:
    """Given entries from report_below_element_captions, move each caption
    (including any contiguous same-style continuation paragraphs) to sit
    immediately before its figure/table.

    Operates at the XML element level. The caption's w:p element(s) are
    detached from their current location and inserted just before the
    preceding w:p (image-bearing) or w:tbl element, preserving original
    ordering of the caption paragraphs.

    Handles multi-paragraph captions: if the target caption is directly
    followed by more paragraphs in the same style, they are moved as a
    group so the continuation stays with the head.

    Indices in the report are interpreted against the CURRENT document state,
    so callers should pass a freshly-generated report (do not cache across
    swaps). Returns the number of caption GROUPS actually moved.
    """
    moved = 0
    tag_p, tag_tbl = qn("w:p"), qn("w:tbl")
    for entry in report:
        # Re-resolve the caption paragraph each iteration since previous swaps
        # change paragraph indices. Match by the preview text to stay robust.
        preview = entry.get("preview", "")
        kind = entry["kind"]
        target_p = None
        for p in doc.paragraphs:
            if p.style is None:
                continue
            sn = p.style.name
            if kind == "figure" and sn != "TME Figure Caption":
                continue
            if kind == "table" and sn != "TME Table Caption":
                continue
            if p.text.strip().startswith(preview.strip()):
                target_p = p
                break
        if target_p is None:
            continue

        first_cap_el = target_p._p

        # Collect any contiguous same-style caption paragraphs following the
        # target. This handles multi-paragraph captions — e.g., a caption body
        # that the author broke across two paragraphs.
        target_pPr = first_cap_el.find(qn("w:pPr"))
        target_pStyle = target_pPr.find(qn("w:pStyle")) if target_pPr is not None else None
        target_style_id = target_pStyle.get(qn("w:val")) if target_pStyle is not None else None

        caption_els = [first_cap_el]
        cursor = first_cap_el.getnext()
        while cursor is not None and cursor.tag == tag_p:
            pPr = cursor.find(qn("w:pPr"))
            pStyle = pPr.find(qn("w:pStyle")) if pPr is not None else None
            style_id = pStyle.get(qn("w:val")) if pStyle is not None else None
            if style_id is not None and style_id == target_style_id:
                caption_els.append(cursor)
                cursor = cursor.getnext()
            else:
                break

        parent = first_cap_el.getparent()
        # Find the preceding image paragraph or table
        prev = first_cap_el.getprevious()
        # Skip empty-paragraph spacers (matching report_below_element_captions
        # semantics — spacers are paragraphs with no text, no w:t, no drawing,
        # no pict)
        while prev is not None and prev.tag == tag_p:
            has_content = (
                prev.find(".//" + qn("w:t")) is not None or
                prev.find(".//" + qn("w:drawing")) is not None or
                prev.find(".//" + qn("w:pict")) is not None
            )
            if has_content:
                break
            prev = prev.getprevious()
        if prev is None:
            continue
        if kind == "figure" and prev.tag != tag_p:
            continue
        if kind == "table" and prev.tag != tag_tbl:
            continue

        for cap_el in caption_els:
            parent.remove(cap_el)
        for cap_el in caption_els:
            prev.addprevious(cap_el)
        moved += 1
    return moved


def normalize_table_cells(doc, skip_indices=(0, 1)) -> int:
    """For content tables (skipping masthead + author card), strip run-level
    font name and size overrides inside cells so content renders at the body
    font (Georgia). Preserves bold/italic which are used for emphasis."""
    n = 0
    for i, table in enumerate(doc.tables):
        if i in skip_indices:
            continue
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        rPr = r._r.find(qn("w:rPr"))
                        if rPr is None:
                            continue
                        for tag in ("w:rFonts", "w:sz", "w:szCs"):
                            for el in rPr.findall(qn(tag)):
                                rPr.remove(el)
                                n += 1
    return n


# ---------- table fixes ----------

def _set_trPr_flag(tr, tag_name: str) -> None:
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    existing = trPr.find(qn(f"w:{tag_name}"))
    if existing is None:
        el = OxmlElement(f"w:{tag_name}")
        trPr.append(el)


def fix_content_tables(doc, skip_indices=(0, 1)) -> int:
    """Center content tables, prevent row splits, repeat header row.

    skip_indices are masthead (0) and the author-card table (1) in our
    standard starter layout. If the docx structure differs, the caller can
    override.
    """
    n = 0
    for i, table in enumerate(doc.tables):
        if i in skip_indices:
            continue
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        tblPr = table._tbl.find(qn("w:tblPr"))
        if tblPr is not None:
            tblInd = tblPr.find(qn("w:tblInd"))
            if tblInd is not None:
                tblInd.set(qn("w:w"), "0")
            jc = tblPr.find(qn("w:jc"))
            if jc is None:
                jc = OxmlElement("w:jc")
                tblPr.append(jc)
            jc.set(qn("w:val"), "center")

        rows = table._tbl.findall(qn("w:tr"))
        for tr in rows:
            _set_trPr_flag(tr, "cantSplit")
        if rows:
            _set_trPr_flag(rows[0], "tblHeader")
        n += 1
    return n


def fix_masthead_grid(doc) -> bool:
    """Defensive: rewrite the masthead table grid in case Word merged it with
    the tagline table (see fixup_moore.py for the original diagnosis)."""
    if not doc.tables:
        return False
    table = doc.tables[0]
    tbl = table._tbl

    BLEED = 90  # twips
    COL0, COL1 = 4651, 7589 + BLEED
    TOTAL = COL0 + COL1

    existing_grid = tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    tblGrid = OxmlElement("w:tblGrid")
    for w in (COL0, COL1):
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        tblGrid.append(col)
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl.insert(0, tblGrid)

    rows = tbl.findall(qn("w:tr"))
    for tr in rows:
        trPr = tr.find(qn("w:trPr"))
        if trPr is not None:
            for tag in ("w:gridBefore", "w:gridAfter", "w:wBefore", "w:wAfter"):
                for el in trPr.findall(qn(tag)):
                    trPr.remove(el)

    if len(rows) >= 1:
        cells = rows[0].findall(qn("w:tc"))
        widths = [COL0, COL1]
        for cell, w in zip(cells, widths):
            tcPr = cell.find(qn("w:tcPr"))
            if tcPr is None:
                continue
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.insert(0, tcW)
            tcW.set(qn("w:w"), str(w))
            tcW.set(qn("w:type"), "dxa")
            gs = tcPr.find(qn("w:gridSpan"))
            if gs is not None:
                tcPr.remove(gs)
    if len(rows) >= 2:
        cells = rows[1].findall(qn("w:tc"))
        if cells:
            cell = cells[0]
            tcPr = cell.find(qn("w:tcPr"))
            if tcPr is not None:
                tcW = tcPr.find(qn("w:tcW"))
                if tcW is None:
                    tcW = OxmlElement("w:tcW")
                    tcPr.insert(0, tcW)
                tcW.set(qn("w:w"), str(TOTAL))
                tcW.set(qn("w:type"), "dxa")
                gs = tcPr.find(qn("w:gridSpan"))
                if gs is None:
                    gs = OxmlElement("w:gridSpan")
                    tcPr.append(gs)
                gs.set(qn("w:val"), "2")

    if tblPr is not None:
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.insert(0, tblW)
        tblW.set(qn("w:w"), str(TOTAL))
        tblW.set(qn("w:type"), "dxa")
    return True


# ---------- footnote font + size fix (zip-level) ----------

def fix_footnote_fonts(docx_path: Path) -> dict:
    src = str(docx_path)
    with zipfile.ZipFile(src, "r") as z:
        if "word/footnotes.xml" not in z.namelist():
            return {"rfonts_rewritten": 0, "rfonts_injected": 0, "sz_stripped": 0}
        with z.open("word/footnotes.xml") as f:
            xml = f.read().decode("utf-8")

    rewritten = 0
    def _rewrite_rfonts(m):
        nonlocal rewritten
        rewritten += 1
        return '<w:rFonts w:ascii="Georgia" w:hAnsi="Georgia" w:cs="Georgia"/>'
    xml = re.sub(r"<w:rFonts\b[^/]*/>", _rewrite_rfonts, xml)

    stripped = 0
    def _count_sub(pattern, s):
        nonlocal stripped
        s2, n = re.subn(pattern, "", s)
        stripped += n
        return s2
    xml = _count_sub(r"<w:sz\b[^/]*/>", xml)
    xml = _count_sub(r"<w:szCs\b[^/]*/>", xml)

    injected = 0
    def _inject_rfonts(m):
        nonlocal injected
        inner = m.group(1)
        if "<w:rFonts" in inner:
            return m.group(0)
        injected += 1
        return f'<w:rPr>{inner}<w:rFonts w:ascii="Georgia" w:hAnsi="Georgia" w:cs="Georgia"/></w:rPr>'
    xml = re.sub(r"<w:rPr>(.*?)</w:rPr>", _inject_rfonts, xml, flags=re.DOTALL)

    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(tmp_fd)
    with zipfile.ZipFile(src, "r") as zin, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/footnotes.xml":
                data = xml.encode("utf-8")
            zout.writestr(item, data)
    shutil.move(tmp_path, src)
    return {"rfonts_rewritten": rewritten, "rfonts_injected": injected, "sz_stripped": stripped}


# ---------- main entry ----------

def run_fixup(docx_path: str) -> dict:
    """Apply the full fixup battery to the docx at docx_path, in place. Returns
    a stats dict for UI display."""
    doc = Document(docx_path)

    update_styles(doc)
    bq_count = remap_block_quotes(doc)
    rescued_refs = rescue_misclassified_references(doc)
    caption_stats = fix_caption_classifications(doc)
    list_n = clear_list_direct_spacing(doc)
    t_count = fix_content_tables(doc)
    masthead_ok = fix_masthead_grid(doc)
    ref_stripped = strip_reference_run_formatting(doc)
    # Run AFTER reclassification so stripping applies to final-style paragraphs
    direct_strip = strip_direct_formatting(doc)
    cell_strip = normalize_table_cells(doc)

    below = report_below_element_captions(doc)

    doc.save(docx_path)

    # Zip-level footnote fix must be after python-docx save.
    fn_stats = fix_footnote_fonts(Path(docx_path))

    return {
        "block_quotes_remapped": bq_count,
        "refs_rescued": rescued_refs,
        "captions": caption_stats,
        "captions_below_element": below,
        "lists_cleared": list_n,
        "tables_centered": t_count,
        "masthead_ok": masthead_ok,
        "references_stripped": ref_stripped,
        "direct_formatting": direct_strip,
        "table_cells_normalized": cell_strip,
        "footnotes": fn_stats,
    }
