"""Parameterized version of moore_build/src/apply_tme_styles.py.

Applies TME paragraph styles to a populated article starter. Detects the body
section by looking for the paragraph containing the Section 1→2 page break
(our cover-builder's structure), removes duplicated cover content pasted in
from the editor's manuscript (title, abstract, author info), and classifies
the remaining paragraphs into TME Body / TME H1 / TME Figure Caption / etc.
"""
from __future__ import annotations

import re
from typing import List, Optional

from docx import Document
from docx.oxml.ns import qn


FIG_PAT = re.compile(r"^\s*Figure\s", re.I)
TAB_PAT = re.compile(r"^\s*Table\s", re.I)
# Matches "LastName, F." or "LastName, F. M." or "LastName, F., &" — reference entry openers
REF_PAT = re.compile(
    r"^[A-ZÀ-ÖØ-Ý][\w'’\-]+,\s+[A-Z]\."
)


def _first_nonempty_run(p):
    for r in p.runs:
        if r.text and r.text.strip():
            return r
    return None


def _find_body_start_index(paragraphs) -> Optional[int]:
    """The body section is the LAST section in our starter structure. Its
    content begins immediately after the last paragraph-embedded sectPr.
    (The very last section's sectPr lives outside any paragraph at the body
    element's tail, so iterating paragraphs naturally stops at the second-to-last
    section's sectPr — which is exactly the one preceding body content.)"""
    last_idx = None
    for i, p in enumerate(paragraphs):
        pPr = p._p.find(qn("w:pPr"))
        if pPr is None:
            continue
        if pPr.find(qn("w:sectPr")) is not None:
            last_idx = i
    return (last_idx + 1) if last_idx is not None else None


def _looks_like_cover_duplicate(text: str, meta) -> bool:
    """True if a pasted paragraph seems to duplicate cover content the app
    already rendered (title, abstract opener, author names, dates, keywords,
    affiliations). Called only on paragraphs at the *top* of the body section.
    """
    t = text.strip()
    if not t:
        return False
    tl = t.lower()
    title_l = (meta.title or "").strip().lower()
    if title_l and (title_l in tl or tl in title_l or tl.startswith(title_l[:40])):
        return True
    # Abstract content: match if the paragraph starts with same first ~40 chars
    abs_opener = (meta.abstract or "").strip()[:40].lower()
    if abs_opener and abs_opener in tl:
        return True
    # Section-header artifacts
    if tl in {"abstract", "keywords"}:
        return True
    if tl.startswith("keywords:") or tl.startswith("keywords "):
        return True
    # Author names, any
    for a in meta.authors or []:
        if a.name and a.name.strip().lower() in tl and len(tl) < 200:
            return True
    # Date lines
    for kw in ("received", "revised", "accepted", "published"):
        if tl.startswith(kw):
            return True
    # Affiliations
    for aff in meta.affiliations or []:
        head = aff.strip().lower()[:25]
        if head and head in tl and len(tl) < 200:
            return True
    # Corresponding author marker
    if "corresponding author" in tl or "corresponding:" in tl:
        return True
    return False


def _style(doc, name: str):
    return doc.styles[name]


def _assign(p, doc, style_name: str, stats: dict) -> None:
    p.style = doc.styles[style_name]
    stats["applied"][style_name] = stats["applied"].get(style_name, 0) + 1


# Mapping from Gemini classifier label → TME style name
_LABEL_TO_STYLE = {
    "heading_1": "TME H1",
    "heading_2": "TME H2",
    "heading_3": "TME H3",
    "body": "TME Body",
    "caption_figure": "TME Figure Caption",
    "caption_table": "TME Table Caption",
    "reference": "TME Reference",
    "block_quote": "TME Block Quote",
    "list_item": "List Paragraph",
    # "skip" left intentionally absent — don't restyle
}


def _heuristic_classify(t: str, p) -> str:
    """Fallback classifier when Gemini is unavailable. Same logic as the
    previous heuristic version — conservative toward 'body' on ambiguity."""
    if REF_PAT.match(t):
        return "reference"
    if FIG_PAT.match(t) and len(t) < 400:
        return "caption_figure"
    if TAB_PAT.match(t) and len(t) < 400:
        return "caption_table"
    first = _first_nonempty_run(p)
    is_bold = bool(first and first.bold)
    is_short = len(t) < 150
    is_labeled_body = ": " in t and len(t.split(": ", 1)[1].strip()) > 20
    has_fill_in = re.search(r"[_—–]{3,}", t) is not None
    tr = t.rstrip()
    ends_sentence = tr.endswith(("?", "!", ":", ",", ";"))
    ends_with_period = tr.endswith(".")
    if (is_bold and is_short
            and not is_labeled_body
            and not has_fill_in
            and not ends_sentence
            and not (ends_with_period and len(t) > 40)):
        return "heading_1"
    return "body"


def apply_styles(docx_path: str, meta) -> dict:
    """Open docx at docx_path, apply TME styles, save in place. Returns stats.

    `meta` is the ArticleMeta that was used to build this starter. Used to
    identify and remove pasted-in duplicates of the cover content, and to
    give the Gemini classifier article context.
    """
    doc = Document(docx_path)
    paragraphs = list(doc.paragraphs)

    body_start = _find_body_start_index(paragraphs)
    if body_start is None:
        # Fallback: start from the paragraph after the last TME Title paragraph.
        last_title_idx = None
        for i, p in enumerate(paragraphs):
            if p.style and p.style.name == "TME Title":
                last_title_idx = i
        body_start = (last_title_idx or 0) + 1

    # Remove leading duplicates of cover content and any leftover placeholder
    deleted_preamble = 0
    placeholder_token = "paste article body here"
    while body_start < len(paragraphs):
        p = paragraphs[body_start]
        t = p.text.strip()
        tl = t.lower()
        is_placeholder = placeholder_token in tl
        if not t or is_placeholder or _looks_like_cover_duplicate(t, meta):
            p._element.getparent().remove(p._element)
            deleted_preamble += 1
            # Refresh paragraphs list since we mutated
            paragraphs = list(doc.paragraphs)
            # body_start index refers to the next element naturally now
            continue
        break

    stats = {
        "deleted_preamble": deleted_preamble,
        "skipped_empty": 0,
        "classifier": "heuristic",  # overwritten to 'gemini' if that path runs
        "applied": {},  # style name → count
    }

    # First pass: handle paragraphs whose SOURCE style already tells us what
    # they are. This short-circuits before calling Gemini and is cheap insurance
    # against Gemini misclassifying things the source already marked.
    body_paragraphs = paragraphs[body_start:]
    pending = []  # (index_in_body_paragraphs, paragraph) for Gemini

    for i, p in enumerate(body_paragraphs):
        t = p.text.strip()
        if not t:
            stats["skipped_empty"] += 1
            continue
        src_style = p.style.name if p.style else ""

        if src_style in ("TMEReference", "TME Reference",
                         "EndNoteBibliography", "EndNote Bibliography"):
            _assign(p, doc, "TME Reference", stats)
            p.paragraph_format.left_indent = None
            p.paragraph_format.first_line_indent = None
            continue
        if src_style == "Caption" or src_style.startswith("Caption"):
            if TAB_PAT.match(t):
                _assign(p, doc, "TME Table Caption", stats)
            else:
                _assign(p, doc, "TME Figure Caption", stats)
            continue
        if src_style == "EndNoteBibliographyTitle":
            _assign(p, doc, "TME H1", stats)
            continue
        if src_style in ("List Paragraph", "ListParagraph"):
            # Leave as List Paragraph; style is already registered.
            continue
        if src_style in ("Normal", "", "Default Paragraph Font", "Body Text"):
            pending.append((i, p))
            continue
        # Unknown source style — let Gemini look at it too
        pending.append((i, p))

    # Second pass: ask Gemini to classify everything that's left
    pending_texts = [p.text.strip() for _, p in pending]
    labels = None
    if pending_texts:
        try:
            from classifier import classify_paragraphs
            labels = classify_paragraphs(
                pending_texts,
                title=meta.title or "",
                abstract=meta.abstract or "",
            )
            stats["classifier"] = "gemini"
        except Exception as e:
            stats["classifier"] = f"heuristic (gemini error: {e})"
            labels = None

    # Apply classifications (Gemini if we have them, heuristic otherwise)
    for idx, (_, p) in enumerate(pending):
        t = p.text.strip()
        if labels is not None:
            label = labels[idx]
        else:
            label = _heuristic_classify(t, p)
        style_name = _LABEL_TO_STYLE.get(label)
        if style_name is None:  # "skip" or unrecognized
            continue
        _assign(p, doc, style_name, stats)
        if style_name == "TME Reference":
            p.paragraph_format.left_indent = None
            p.paragraph_format.first_line_indent = None

    # Re-establish the body section's running footer + headers. Word silently
    # relinks a newly-populated section's footer to the previous section when
    # the editor pastes into it; calling our footer builders again resets this.
    try:
        from tme_template.headers_footers import set_running_footer, set_running_headers
        body_section = doc.sections[-1]
        cite_last_names = " & ".join(
            a.name.rsplit(" ", 1)[-1] for a in (meta.authors or []) if a.name
        ) or "Author"
        short_title = meta.title if len(meta.title) < 60 else meta.title[:57] + "..."
        set_running_headers(
            doc, author_cite=cite_last_names, short_title=short_title,
            section=body_section,
        )
        set_running_footer(doc, section=body_section)
        stats["footer_restored"] = True
    except Exception as e:
        stats["footer_restored"] = f"failed: {e}"

    doc.save(docx_path)
    return stats
