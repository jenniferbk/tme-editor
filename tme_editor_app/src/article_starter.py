"""Generalized article-starter builder: turns ArticleMeta + headshots into a
formatted TME starter .docx with cover page + placeholder body section.

This is the generalized version of moore_pipeline.moore_starter."""
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.shared import Inches

from tme_template.masthead import MastheadData, add_masthead
from tme_template.tagline import add_tagline_strip
from tme_template.cover_page import AuthorEntry, CoverData, add_research_article_cover
from tme_template.cover_footer import add_cover_footer
from tme_template.styles import (
    register_body_style, register_title_style,
    register_heading_styles, register_remaining_styles,
)
from tme_template.page_setup import configure_page_setup, configure_zero_margins
from tme_template.oxml_helpers import add_section_break_next_page, add_continuous_section_break
from tme_template.headers_footers import set_running_headers, set_running_footer


REPO_ROOT = Path(__file__).resolve().parents[2]
LOGO_H = str(REPO_ROOT / "assets" / "tme-logo.jpg")


PLACEHOLDER_TEXT = (
    "[Paste article body here — apply TME Body style to paragraphs, "
    "TME H1/H2/H3 to headings, TME Figure Caption to captions, "
    "TME Reference to reference entries.]"
)


def _format_citation(meta) -> str:
    """APA-ish one-line citation suitable for the cover footer."""
    if not meta.authors:
        authors = "Anonymous"
    else:
        parts = []
        for a in meta.authors:
            last = a.name.rsplit(" ", 1)[-1]
            initials = " ".join(n[0] + "." for n in a.name.split(" ")[:-1] if n)
            parts.append(f"{last}, {initials}".strip(", "))
        if len(parts) == 1:
            authors = parts[0]
        elif len(parts) == 2:
            authors = " & ".join(parts)
        else:
            authors = ", ".join(parts[:-1]) + f", & {parts[-1]}"
    title_short = meta.title if len(meta.title) < 120 else meta.title[:117] + "..."
    return (
        f"{authors} ({meta.year}). {title_short}. "
        f"The Mathematics Educator, {meta.volume}({meta.number}), "
        f"{meta.pages.replace('–', '–')}."
    )


def build_article_starter(*, meta, headshots: Dict[str, Path], out_path: Path) -> None:
    """Generate a starter .docx for any article. `headshots` maps author name to
    a framed headshot path (as produced by prepare_headshots)."""
    doc = Document()
    configure_page_setup(doc)
    register_body_style(doc)
    register_title_style(doc)
    register_heading_styles(doc)
    register_remaining_styles(doc)

    masthead_section = doc.sections[0]
    configure_zero_margins(masthead_section)
    masthead_section.bottom_margin = Inches(0.5)
    add_masthead(doc, MastheadData(
        article_type=meta.article_type,
        volume=meta.volume, number=meta.number, year=meta.year, pages=meta.pages,
        doi=meta.doi or None,
        issn_print="1062-9017", issn_online="2331-4451",
        logo_path=LOGO_H,
    ))
    add_tagline_strip(doc)
    add_cover_footer(
        masthead_section,
        citation=_format_citation(meta),
    )

    cover_body_section = add_continuous_section_break(doc)
    cover_body_section.page_width = Inches(8.5)
    cover_body_section.page_height = Inches(11)
    cover_body_section.top_margin = Inches(0.3)
    cover_body_section.bottom_margin = Inches(0.3)
    cover_body_section.left_margin = Inches(0.5)
    cover_body_section.right_margin = Inches(0.5)
    cover_body_section.footer.is_linked_to_previous = False
    for p in list(cover_body_section.footer.paragraphs):
        p._p.getparent().remove(p._p)

    author_entries = []
    for a in meta.authors:
        headshot_path = headshots.get(a.name)
        author_entries.append(AuthorEntry(
            name=a.name,
            affiliation_num=a.affiliation_num,
            role=a.role,
            bio=a.bio,
            headshot_path=str(headshot_path) if headshot_path else None,
            corresponding=a.corresponding,
            email=a.email,
        ))
    add_research_article_cover(doc, CoverData(
        title=meta.title,
        authors=author_entries,
        affiliations=meta.affiliations,
        dates={
            "Received": meta.received or "TBD",
            "Revised": meta.revised or "TBD",
            "Accepted": meta.accepted or "TBD",
            "Published": meta.published or f"{meta.year}",
        },
        abstract=meta.abstract,
        keywords=meta.keywords,
    ))

    body_section = add_section_break_next_page(doc)
    cite_last_names = " & ".join(
        a.name.rsplit(" ", 1)[-1] for a in meta.authors
    ) or "Author"
    short_title = meta.title if len(meta.title) < 60 else meta.title[:57] + "..."
    set_running_headers(doc,
        author_cite=cite_last_names,
        short_title=short_title,
        section=body_section)
    set_running_footer(doc, section=body_section)

    doc.add_paragraph(PLACEHOLDER_TEXT, style="TME Body")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
