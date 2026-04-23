"""Tests for the new caption-below-element detection and swap helpers."""
from docx import Document

import fixup
from tme_template.styles import (
    register_body_style, register_title_style,
    register_heading_styles, register_remaining_styles,
)


def _make_doc_with_styles():
    doc = Document()
    register_body_style(doc)
    register_title_style(doc)
    register_heading_styles(doc)
    register_remaining_styles(doc)
    return doc


def _add_fake_drawing_paragraph(doc):
    """Add a paragraph whose XML contains a w:drawing element so the detector
    treats it as an image-bearing paragraph without needing real image bytes."""
    p = doc.add_paragraph("", style="TME Body")
    from docx.oxml import OxmlElement
    run = p.add_run()
    drawing = OxmlElement("w:drawing")
    run._r.append(drawing)
    return p


def test_report_flags_figure_caption_below_drawing():
    doc = _make_doc_with_styles()
    # figure-above-caption pattern (APA-correct): no report
    doc.add_paragraph("Figure 1. A thing.", style="TME Figure Caption")
    _add_fake_drawing_paragraph(doc)
    # caption-below-drawing pattern (flag this)
    _add_fake_drawing_paragraph(doc)
    doc.add_paragraph("Figure 2. Another thing.", style="TME Figure Caption")
    result = fixup.report_below_element_captions(doc)
    kinds = [r["kind"] for r in result]
    assert kinds == ["figure"]
    assert "Figure 2" in result[0]["preview"]


def test_report_flags_table_caption_below_table():
    doc = _make_doc_with_styles()
    # caption-above-table (correct): no report
    doc.add_paragraph("Table 1. Correct position.", style="TME Table Caption")
    doc.add_table(rows=2, cols=2)
    # table-above-caption (flag)
    doc.add_table(rows=2, cols=2)
    doc.add_paragraph("Table 2. Wrong position.", style="TME Table Caption")
    result = fixup.report_below_element_captions(doc)
    kinds = [r["kind"] for r in result]
    assert kinds == ["table"]
    assert "Table 2" in result[0]["preview"]


def test_run_fixup_returns_captions_below_element_key():
    doc = _make_doc_with_styles()
    _add_fake_drawing_paragraph(doc)
    doc.add_paragraph("Figure 1. Below.", style="TME Figure Caption")
    # Save to a temp path, run fixup, check the stats dict
    import tempfile
    from pathlib import Path
    with tempfile.TemporaryDirectory() as td:
        path = Path(td) / "t.docx"
        doc.save(str(path))
        stats = fixup.run_fixup(str(path))
    assert "captions_below_element" in stats
    assert isinstance(stats["captions_below_element"], list)


def test_swap_captions_above_moves_figure_caption_up():
    doc = _make_doc_with_styles()
    _add_fake_drawing_paragraph(doc)
    doc.add_paragraph("Figure 1. Below.", style="TME Figure Caption")
    doc.add_paragraph("Some body paragraph.", style="TME Body")

    report = fixup.report_below_element_captions(doc)
    assert len(report) == 1

    moved = fixup.swap_captions_above(doc, report)
    assert moved == 1

    # After swap, caption should precede the drawing paragraph
    texts = [p.text for p in doc.paragraphs]
    cap_idx = next(i for i, t in enumerate(texts) if t.startswith("Figure 1"))
    # Drawing paragraph is the one right AFTER the caption now
    assert cap_idx < len(texts) - 1
    # Re-run report: should be empty now
    report2 = fixup.report_below_element_captions(doc)
    assert report2 == []


def test_swap_captions_above_moves_table_caption_up():
    doc = _make_doc_with_styles()
    doc.add_table(rows=2, cols=2)
    doc.add_paragraph("Table 1. Below.", style="TME Table Caption")

    report = fixup.report_below_element_captions(doc)
    assert len(report) == 1
    moved = fixup.swap_captions_above(doc, report)
    assert moved == 1
    assert fixup.report_below_element_captions(doc) == []
