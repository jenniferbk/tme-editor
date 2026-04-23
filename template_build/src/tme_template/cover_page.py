"""Research-article cover page layout: title, authors, abstract, author block."""
from dataclasses import dataclass, field
from typing import List, Optional, Dict

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from tme_template.colors import LINE, META, UGA_RED
from tme_template.oxml_helpers import apply_bottom_rule, remove_cell_borders, set_cell_margins


@dataclass
class AuthorEntry:
    name: str
    affiliation_num: int
    role: Optional[str]
    bio: str
    headshot_path: Optional[str]
    corresponding: bool = False
    email: Optional[str] = None


@dataclass
class CoverData:
    title: str
    authors: List[AuthorEntry]
    affiliations: List[str]
    dates: Dict[str, str]  # ordered: Received, Revised, Accepted, Published
    abstract: str
    keywords: List[str]


def _red_label(paragraph, text: str, size_pt: float = 9):
    r = paragraph.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size_pt)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xBA, 0x0C, 0x2F)
    return r


def add_research_article_cover(doc, data: CoverData) -> None:
    # Title
    p = doc.add_paragraph(style="TME Title")
    p.add_run(data.title)

    # Authors line — built run by run so affiliation numbers are true superscripts
    authors_p = doc.add_paragraph()
    authors_p.paragraph_format.space_before = Pt(0)
    authors_p.paragraph_format.space_after = Pt(3)
    for i, a in enumerate(data.authors):
        if i == 0:
            pass  # first author — no separator
        elif i == len(data.authors) - 1:
            sep = authors_p.add_run(", and " if len(data.authors) > 2 else " and ")
            sep.font.name = "Georgia"
            sep.font.size = Pt(11)
        else:
            sep = authors_p.add_run(", ")
            sep.font.name = "Georgia"
            sep.font.size = Pt(11)
        r_name = authors_p.add_run(a.name)
        r_name.font.name = "Georgia"
        r_name.font.size = Pt(11)
        r_aff = authors_p.add_run(str(a.affiliation_num))
        r_aff.font.name = "Georgia"
        r_aff.font.size = Pt(11)
        r_aff.font.superscript = True
        if a.corresponding:
            r_dag = authors_p.add_run("†")
            r_dag.font.name = "Georgia"
            r_dag.font.size = Pt(11)
            r_dag.font.superscript = True
            r_dag.font.color.rgb = RGBColor(0xBA, 0x0C, 0x2F)
            r_dag.font.bold = True

    # Affiliations — superscript numbers, no brackets
    for i, aff in enumerate(data.affiliations, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r_num = p.add_run(str(i))
        r_num.font.name = "Georgia"
        r_num.font.size = Pt(9)
        r_num.font.superscript = True
        r_num.font.color.rgb = RGBColor.from_string(META)
        r_txt = p.add_run(aff)
        r_txt.font.name = "Georgia"
        r_txt.font.size = Pt(9)
        r_txt.font.color.rgb = RGBColor.from_string(META)

    # Corresponding author line
    corres = next((a for a in data.authors if a.corresponding), None)
    if corres and corres.email:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(0)
        r_dag = p.add_run("† ")
        r_dag.font.color.rgb = RGBColor(0xBA, 0x0C, 0x2F)
        r_dag.font.bold = True
        r = p.add_run(f"Corresponding author: {corres.email}")
        r.font.name = "Georgia"
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = RGBColor.from_string(META)

    # Dates row
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    for i, (label, date) in enumerate(data.dates.items()):
        if i > 0:
            sep = p.add_run("  ·  ")
            sep.font.name = "Arial"
            sep.font.size = Pt(9)
            sep.font.color.rgb = RGBColor.from_string(LINE)
        r_l = p.add_run(label.upper())
        r_l.font.name = "Arial"
        r_l.font.size = Pt(9)
        r_l.font.bold = True
        r_l.font.color.rgb = RGBColor.from_string(META)
        r_d = p.add_run(f" {date}")
        r_d.font.name = "Arial"
        r_d.font.size = Pt(9)
        r_d.font.color.rgb = RGBColor.from_string(META)

    # Rule
    rule_p = doc.add_paragraph()
    rule_p.paragraph_format.space_before = Pt(0)
    rule_p.paragraph_format.space_after = Pt(8)
    apply_bottom_rule(rule_p, hex_color=LINE, width_pt=1)

    # About the Authors label
    ab_lbl = doc.add_paragraph()
    ab_lbl.paragraph_format.space_before = Pt(0)
    ab_lbl.paragraph_format.space_after = Pt(4)
    _red_label(ab_lbl, "ABOUT THE AUTHORS")

    # Author block — 3-column table (one cell per author)
    n = len(data.authors)
    col_width = Inches(7.5 / n)
    tbl = doc.add_table(rows=1, cols=n)
    tbl.style = "Table Grid"

    # Keep author row from breaking across pages
    row = tbl.rows[0]
    trPr = row._tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)

    for col_idx, a in enumerate(data.authors):
        cell = tbl.cell(0, col_idx)
        remove_cell_borders(cell)
        set_cell_margins(cell, top=0, bottom=0, left=80, right=80)
        # Set cell width
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = tcPr.find(qn("w:tcW"))
        if tcW is None:
            tcW = OxmlElement("w:tcW")
            tcPr.append(tcW)
        tcW.set(qn("w:w"), str(int(col_width.pt * 20)))  # twentieths of a point
        tcW.set(qn("w:type"), "dxa")

        # Paragraph 1: headshot (centered)
        p_img = cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(2)
        p_img.paragraph_format.space_after = Pt(2)
        if a.headshot_path:
            img_run = p_img.add_run()
            img_run.add_picture(a.headshot_path, height=Inches(0.65))
        # else leave paragraph empty

        # Paragraph 2: author name (italic bold Georgia 9.5pt)
        p_name = cell.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_name.paragraph_format.space_before = Pt(2)
        p_name.paragraph_format.space_after = Pt(2)
        r_name = p_name.add_run(a.name)
        r_name.font.name = "Georgia"
        r_name.font.size = Pt(9.5)
        r_name.font.bold = True
        r_name.font.italic = True

        # Paragraph 3: bio (Georgia 8.5pt, justified, 1.2 line spacing)
        p_bio = cell.add_paragraph()
        p_bio.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_bio.paragraph_format.space_before = Pt(2)
        p_bio.paragraph_format.space_after = Pt(2)
        p_bio.paragraph_format.line_spacing = 1.2
        r_bio = p_bio.add_run(a.bio)
        r_bio.font.name = "Georgia"
        r_bio.font.size = Pt(8.5)

    # Rule separating author block from abstract
    rule_p2 = doc.add_paragraph()
    rule_p2.paragraph_format.space_before = Pt(0)
    rule_p2.paragraph_format.space_after = Pt(8)
    apply_bottom_rule(rule_p2, hex_color=LINE, width_pt=1)

    # Abstract label and text
    lbl = doc.add_paragraph()
    lbl.paragraph_format.space_before = Pt(0)
    lbl.paragraph_format.space_after = Pt(4)
    _red_label(lbl, "ABSTRACT")

    # Abstract — explicit compact formatting (Georgia 10pt, 1.3 line spacing, justified)
    ab = doc.add_paragraph()
    ab.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ab.paragraph_format.space_before = Pt(0)
    ab.paragraph_format.space_after = Pt(6)
    ab.paragraph_format.line_spacing = 1.3
    ab_run = ab.add_run(data.abstract)
    ab_run.font.name = "Georgia"
    ab_run.font.size = Pt(10)

    # Keywords
    kw = doc.add_paragraph()
    kw.paragraph_format.space_before = Pt(0)
    kw.paragraph_format.space_after = Pt(0)
    r1 = kw.add_run("Keywords: ")
    r1.bold = True
    r1.font.name = "Georgia"
    r1.font.size = Pt(9.5)
    r_kw = kw.add_run(" · ".join(data.keywords))
    r_kw.font.name = "Georgia"
    r_kw.font.size = Pt(9.5)
