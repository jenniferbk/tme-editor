"""Front-matter page generators: issue cover, editorial staff, formal title page."""
import io
from dataclasses import dataclass, field
from typing import Dict, List, Optional

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from tme_template.colors import INK, META, UGA_RED
from tme_template.oxml_helpers import apply_bottom_rule, remove_cell_borders, set_cell_shading


@dataclass
class IssueInfo:
    volume: int
    number: int
    year: int
    season: str  # "Spring", "Summer", "Fall", "Winter"
    cover_artist: Optional[str]
    portrait_logo_path: str


@dataclass
class StaffRoster:
    editors: List[str]
    associate_editors: List[str]
    advisor: Optional[str]
    copy_editor: Optional[str]
    mesa_officers: Dict[str, str]
    mesa_term: str


def _open_image_as_rgb_stream(path: str) -> io.BytesIO:
    """Open an image file, convert to RGB if needed, return as in-memory JPEG stream.

    python-docx cannot handle CMYK JPEGs directly. This function normalizes
    the image to sRGB before inserting.
    """
    from PIL import Image
    img = Image.open(path)
    if img.mode != "RGB":
        img = img.convert("RGB")
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=95)
    buf.seek(0)
    return buf


def _red_run(p, text, *, size_pt=13, bold=True):
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size_pt)
    r.font.bold = bold
    r.font.color.rgb = RGBColor(0xBA, 0x0C, 0x2F)
    return r


FOOTER_CREAM = "FAFAF7"


def add_issue_cover_page(doc, issue: IssueInfo) -> None:
    """Portrait logo, vol/no in red caps, season/year in italic Georgia, credit at bottom.

    Content is wrapped in a full-page single-cell table with cream fill so only
    this page has the cream background (Word page backgrounds are document-wide).
    """
    # Full-page single-cell table with cream fill, no borders
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    # No fill: issue cover uses white background (Word default) so the
    # portrait logo's white background doesn't clash with a cream fill.
    remove_cell_borders(cell)

    # Set approximate page-fill height (~9" for 11" page with 1" top+bottom margins)
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(Inches(9).pt * 20)))  # twips
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)

    def _add_p():
        """Add a paragraph inside the cell (after the first existing paragraph)."""
        p = cell.add_paragraph()
        return p

    # Use the existing first paragraph as a spacer
    p_spacer = cell.paragraphs[0]
    p_spacer.paragraph_format.space_before = Pt(60)

    p_logo = _add_p()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        img_stream = _open_image_as_rgb_stream(issue.portrait_logo_path)
        p_logo.add_run().add_picture(img_stream, height=Inches(4.0))
    except (FileNotFoundError, OSError) as e:
        import sys
        print(f"Warning: could not load portrait logo: {e}", file=sys.stderr)

    p_vol = _add_p()
    p_vol.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_vol.paragraph_format.space_before = Pt(30)
    _red_run(p_vol, f"VOLUME {issue.volume}  ·  NUMBER {issue.number}", size_pt=13)

    p_season = _add_p()
    p_season.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_season.add_run(f"{issue.season} {issue.year}")
    r.font.name = "Georgia"
    r.font.size = Pt(14)
    r.font.italic = True
    r.font.color.rgb = RGBColor.from_string(META)

    p_spacer2 = _add_p()
    p_spacer2.paragraph_format.space_before = Pt(100)

    p_cred = _add_p()
    p_cred.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_cred.add_run(
        "An Official Publication of the Mathematics Education Student Association"
        "  ·  University of Georgia"
    )
    r1.font.name = "Arial"
    r1.font.size = Pt(10)
    r1.font.italic = True
    r1.font.color.rgb = RGBColor.from_string(META)

    if issue.cover_artist:
        p_art = _add_p()
        p_art.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p_art.add_run(f"Cover art by {issue.cover_artist}")
        r2.font.name = "Arial"
        r2.font.size = Pt(10)
        r2.font.italic = True
        r2.font.color.rgb = RGBColor.from_string(META)


def _section_label(doc, text: str):
    p = doc.add_paragraph()
    _red_run(p, text, size_pt=11)
    return p


def _role_group(doc, role_label: str, names: List[str]):
    p_role = doc.add_paragraph()
    r = p_role.add_run(role_label)
    r.font.name = "Georgia"
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = RGBColor.from_string(META)
    for n in names:
        p_n = doc.add_paragraph()
        rn = p_n.add_run(n)
        rn.font.name = "Georgia"
        rn.font.size = Pt(12.5)


def _section_label_in_cell(cell, text: str):
    p = cell.add_paragraph()
    _red_run(p, text, size_pt=11)
    return p


def _role_group_in_cell(cell, role_label: str, names: List[str]):
    p_role = cell.add_paragraph()
    r = p_role.add_run(role_label)
    r.font.name = "Georgia"
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = RGBColor.from_string(META)
    for n in names:
        p_n = cell.add_paragraph()
        rn = p_n.add_run(n)
        rn.font.name = "Georgia"
        rn.font.size = Pt(12.5)


def add_editorial_staff_page(doc, issue: IssueInfo, roster: StaffRoster) -> None:
    """Uses masthead + two-column table layout. Caller must add masthead + tagline first.

    Left column: EDITORIAL group. Right column: MESA OFFICERS group.
    """
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(3.25)
    table.columns[1].width = Inches(3.25)
    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)
    remove_cell_borders(left_cell)
    remove_cell_borders(right_cell)

    # Left column: EDITORIAL group — use existing first paragraph then add_paragraph
    left_cell.paragraphs[0]._p.getparent().remove(left_cell.paragraphs[0]._p)

    _section_label_in_cell(left_cell, "EDITORIAL BOARD")
    _role_group_in_cell(left_cell, "Editors", roster.editors)
    _role_group_in_cell(left_cell, "Associate Editors", roster.associate_editors)
    if roster.advisor:
        _role_group_in_cell(left_cell, "Advisor", [roster.advisor])
    if roster.copy_editor:
        _role_group_in_cell(left_cell, "Copy Editor", [roster.copy_editor])

    # Right column: MESA OFFICERS group
    right_cell.paragraphs[0]._p.getparent().remove(right_cell.paragraphs[0]._p)

    _section_label_in_cell(right_cell, f"MESA OFFICERS {roster.mesa_term}")
    for role, name in roster.mesa_officers.items():
        _role_group_in_cell(right_cell, role, [name])


def add_formal_title_page(doc, issue: IssueInfo) -> None:
    """Pure typography — wordmark + rules + issue info."""
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(120)

    p_wm = doc.add_paragraph()
    p_wm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_wm.add_run("THE MATHEMATICS EDUCATOR")
    r.font.name = "Georgia"
    r.font.size = Pt(36)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(INK)

    # Rule 1 — red bottom border paragraph, constrained to ~200px width
    p_rule = doc.add_paragraph()
    p_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rule.paragraph_format.left_indent = Inches(2.25)
    p_rule.paragraph_format.right_indent = Inches(2.25)
    apply_bottom_rule(p_rule, hex_color="BA0C2F", width_pt=2)

    p_ok = doc.add_paragraph()
    p_ok.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_ok.add_run("An Official Publication of the\nMathematics Education Student Association\nThe University of Georgia")
    r.font.name = "Georgia"
    r.font.size = Pt(13)
    r.font.italic = True
    r.font.color.rgb = RGBColor.from_string(META)

    # Rule 2 — red bottom border paragraph, same width constraint
    p_rule2 = doc.add_paragraph()
    p_rule2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rule2.paragraph_format.left_indent = Inches(2.25)
    p_rule2.paragraph_format.right_indent = Inches(2.25)
    apply_bottom_rule(p_rule2, hex_color="BA0C2F", width_pt=2)

    p_iss = doc.add_paragraph()
    p_iss.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _red_run(p_iss, f"VOLUME {issue.volume}  ·  NUMBER {issue.number}", size_pt=13)

    p_year = doc.add_paragraph()
    p_year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_year.add_run(f"{issue.season} {issue.year}")
    r.font.name = "Georgia"
    r.font.size = Pt(14)
    r.font.italic = True
    r.font.color.rgb = RGBColor.from_string(META)
