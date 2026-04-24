"""Generate the light-gray tagline strip that sits beneath the masthead."""
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from tme_template.colors import LIGHT_PANEL_GRAY, META, UGA_RED
from tme_template.oxml_helpers import (
    force_table_full_width,
    remove_cell_borders,
    set_cell_margins,
    set_cell_shading,
)


TAGLINE = "Cultivating scholarly discourse in mathematics education since 1990"
META_LINE = ("Published by the Mathematics Education Student Association"
             "  ·  University of Georgia  ·  Peer Reviewed  ·  Open Access")


def _red_run(paragraph, text: str, size_pt: float):
    r = paragraph.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size_pt)
    r.font.color.rgb = RGBColor.from_string(UGA_RED)
    return r


def _gray_run(paragraph, text: str, *, name="Georgia", size_pt=11.0, italic=False):
    r = paragraph.add_run(text)
    r.font.name = name
    r.font.size = Pt(size_pt)
    r.font.italic = italic
    r.font.color.rgb = RGBColor.from_string(META)
    return r


def add_tagline_strip(doc) -> None:
    # Bleed to match masthead's right-edge treatment (see masthead.py).
    TOTAL_WIDTH = 8.5
    BLEED_INCHES = 0.063
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(TOTAL_WIDTH + BLEED_INCHES)
    cell = table.cell(0, 0)
    remove_cell_borders(cell)
    set_cell_shading(cell, LIGHT_PANEL_GRAY)
    force_table_full_width(table, total_width_inches=TOTAL_WIDTH + BLEED_INCHES)

    # python-docx's default single-cell width for `add_table(rows=1, cols=1)` is
    # the normal content width (~6.5"), not the page width. Setting the column
    # width above doesn't propagate to the cell's tcW in this case, so the gray
    # fill stops short of the right edge. Set the cell's tcW explicitly to
    # match the full table grid.
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.insert(0, tcW)
    tcW.set(qn("w:w"), str(int((TOTAL_WIDTH + BLEED_INCHES) * 1440)))
    tcW.set(qn("w:type"), "dxa")

    set_cell_margins(cell, top=80, bottom=80, left=160, right=160)

    # Tagline paragraph: ◆ italic-tagline ◆
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after = Pt(2)
    _red_run(p1, "◆ ", size_pt=9.5)
    _gray_run(p1, TAGLINE, size_pt=9.5, italic=True)
    _red_run(p1, " ◆", size_pt=9.5)

    # Meta line
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(4)
    _gray_run(p2, META_LINE, name="Arial", size_pt=8.5)
