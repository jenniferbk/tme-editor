"""The cover-page footer: HOW TO CITE + citation, full-width single cell."""
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from tme_template.colors import FOOTER_CREAM, TEXT_MUTED, UGA_RED
from tme_template.oxml_helpers import (
    force_table_full_width,
    remove_cell_borders,
    set_cell_margins,
    set_cell_shading,
)


def add_cover_footer(section, *, citation: str) -> None:
    """Place the cover footer (HOW TO CITE + citation) inside the given section's
    footer slot. Appears at the bottom of every page in this section.
    Since the cover section is exactly one page (with continuous break above and
    a next-page break after), it will only appear on the cover page."""
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in list(footer.paragraphs):
        p._p.getparent().remove(p._p)

    table = footer.add_table(rows=1, cols=1, width=Inches(7.5))
    table.autofit = False
    table.columns[0].width = Inches(7.5)
    force_table_full_width(table, total_width_inches=7.5, left_indent_inches=0.5)

    cell = table.cell(0, 0)
    remove_cell_borders(cell)
    set_cell_shading(cell, FOOTER_CREAM)
    set_cell_margins(cell, top=120, bottom=120, left=120, right=120)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    r_label = p.add_run("HOW TO CITE  ")
    r_label.font.name = "Arial"
    r_label.font.size = Pt(8)
    r_label.font.bold = True
    r_label.font.color.rgb = RGBColor.from_string(UGA_RED)
    r_cite = p.add_run(citation)
    r_cite.font.name = "Arial"
    r_cite.font.size = Pt(8)
    r_cite.font.color.rgb = RGBColor.from_string(TEXT_MUTED)
