"""Verso/recto running headers and the static running footer."""
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from tme_template.colors import INK, LINE, META
from tme_template.oxml_helpers import apply_top_rule


def _clear(header_or_footer):
    for p in list(header_or_footer.paragraphs):
        p._p.getparent().remove(p._p)


def _add_italic_gray_line(container, text: str, align):
    p = container.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.font.name = "Georgia"
    r.font.size = Pt(10.5)
    r.font.italic = True
    r.font.color.rgb = RGBColor.from_string(META)
    return p


def set_running_headers(doc, author_cite: str, short_title: str, section=None) -> None:
    """Verso (even): author cite flush left. Recto (odd): short title flush right.

    If *section* is provided, operate only on that section object.
    When None (default), operate on all sections (backward-compatible).
    """
    if section is not None:
        section.header.is_linked_to_previous = False
        section.even_page_header.is_linked_to_previous = False
    sections = [section] if section is not None else doc.sections
    for sec in sections:
        # Recto / odd
        recto = sec.header
        _clear(recto)
        _add_italic_gray_line(recto, short_title, WD_ALIGN_PARAGRAPH.RIGHT)
        # Verso / even
        verso = sec.even_page_header
        _clear(verso)
        _add_italic_gray_line(verso, author_cite, WD_ALIGN_PARAGRAPH.LEFT)


def _add_page_field(run):
    """Insert a { PAGE } field into a run."""
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)


def _build_footer(container):
    """Centered page number with a 1pt top rule. Same layout on verso and recto."""
    _clear(container)
    p = container.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_top_rule(p, hex_color=LINE, width_pt=1)
    run = p.add_run()
    run.font.name = "Georgia"
    run.font.size = Pt(9.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(INK)
    _add_page_field(run)


def set_running_footer(doc, section=None) -> None:
    """Body-page footer: centered page number only. No copyright, no license line.

    If *section* is provided, operate only on that section object.
    When None (default), operate on all sections (backward-compatible).
    """
    if section is not None:
        section.footer.is_linked_to_previous = False
        section.even_page_footer.is_linked_to_previous = False
    sections = [section] if section is not None else doc.sections
    for sec in sections:
        _build_footer(sec.footer)
        _build_footer(sec.even_page_footer)
