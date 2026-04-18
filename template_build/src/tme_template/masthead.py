"""Generate the TME masthead — 2-cell table, logo on black left, metadata on UGA red right."""
from dataclasses import dataclass
from typing import Optional

from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from tme_template.colors import BLACK, UGA_RED
from tme_template.oxml_helpers import (
    force_table_full_width,
    remove_cell_borders,
    set_cell_margins,
    set_cell_shading,
)


@dataclass
class MastheadData:
    article_type: str              # e.g., "RESEARCH ARTICLE"
    volume: int
    number: int
    year: int
    pages: Optional[str]           # "1–24" for articles; None for Editorial Staff/TOC
    doi: Optional[str]             # doi.org/... ; None if no DOI
    issn_print: str
    issn_online: str
    logo_path: str


def _set_run_font(run, *, name="Arial", size_pt=11, bold=False, italic=False, white=True):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _vol_line_text(d: MastheadData) -> str:
    base = f"Vol. {d.volume}, No. {d.number} ({d.year})"
    if d.pages:
        base += f", pp. {d.pages}"
    return base


def add_masthead(doc, data: MastheadData) -> None:
    """Append the masthead (2-col table) to a document body."""
    table = doc.add_table(rows=1, cols=2)
    # Fixed widths: left 38%, right 62% of full 8.5" page width (zero margins).
    # Word in Compatibility Mode renders zero-margin content slightly inside
    # the page edge, leaving a visible gap. Extending the right column by a
    # small bleed (~0.063") closes the gap. The logo column stays at 38%.
    TOTAL_WIDTH = 8.5
    BLEED_INCHES = 0.063
    table.autofit = False
    table.columns[0].width = Inches(TOTAL_WIDTH * 0.38)                   # ~3.23"
    table.columns[1].width = Inches(TOTAL_WIDTH * 0.62 + BLEED_INCHES)    # ~5.33"
    table.rows[0].height = Pt(75)

    force_table_full_width(table, total_width_inches=TOTAL_WIDTH + BLEED_INCHES)

    left, right = table.cell(0, 0), table.cell(0, 1)
    for c in (left, right):
        remove_cell_borders(c)
    set_cell_shading(left, BLACK)
    set_cell_shading(right, UGA_RED)

    # Set explicit cell margins after zeroing default table-level margins
    set_cell_margins(left, top=80, bottom=80, left=80, right=80)
    set_cell_margins(right, top=80, bottom=80, left=160, right=160)

    # Left cell: logo image, centered
    left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    try:
        p = left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(data.logo_path, width=Inches(3.0))
    except FileNotFoundError:
        import sys
        print(f"Warning: logo not found at {data.logo_path}", file=sys.stderr)

    # Right cell: stacked metadata, right-aligned
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Clear the default paragraph to have a clean slate
    right.paragraphs[0].text = ""

    def _add_line(text: str, *, size_pt: float, bold: bool = False, letter_spacing_2: bool = False):
        p = right.paragraphs[0] if (len(right.paragraphs) == 1 and right.paragraphs[0].text == "") else right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        _set_run_font(run, size_pt=size_pt, bold=bold)
        return p

    _add_line(data.article_type, size_pt=11, bold=True)
    _add_line(_vol_line_text(data), size_pt=16, bold=True)
    _add_line(data.doi or "", size_pt=10.5)
    _add_line(f"ISSN\u00a0{data.issn_print}\u00a0(print)\u00a0\u00a0{data.issn_online}\u00a0(online)", size_pt=10)
