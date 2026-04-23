"""Register TME paragraph and character styles into a Document."""
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Emu, RGBColor


def _get_or_add_paragraph_style(doc, name: str):
    """Return existing style by name, or add a new paragraph style."""
    if name in [s.name for s in doc.styles]:
        return doc.styles[name]
    return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def register_body_style(doc) -> None:
    style = _get_or_add_paragraph_style(doc, "TME Body")
    style.font.name = "Georgia"
    style.font.size = Pt(11.5)
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.25
    # 1.5em at 11.5pt ≈ 17.25pt ≈ 219075 EMU (12700 EMU per pt)
    pf.first_line_indent = Pt(17.25)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def register_title_style(doc) -> None:
    style = _get_or_add_paragraph_style(doc, "TME Title")
    style.font.name = "Georgia"
    style.font.size = Pt(18)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    pf = style.paragraph_format
    pf.line_spacing = 1.15
    pf.space_before = Pt(12)
    pf.space_after = Pt(10)


def register_heading_styles(doc) -> None:
    """H1 has a red left rule applied at paragraph level via left border;
    see note in apply_red_left_rule helper (added in Task 7)."""
    h1 = _get_or_add_paragraph_style(doc, "TME H1")
    h1.font.name = "Georgia"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.left_indent = Pt(10)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.keep_together = True

    h2 = _get_or_add_paragraph_style(doc, "TME H2")
    h2.font.name = "Georgia"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.italic = True
    h2.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.keep_together = True

    h3 = _get_or_add_paragraph_style(doc, "TME H3")
    h3.font.name = "Georgia"
    h3.font.size = Pt(11.5)
    h3.font.bold = False
    h3.font.italic = True
    h3.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True
    h3.paragraph_format.keep_together = True


def register_remaining_styles(doc) -> None:
    # Figure captions go BELOW figures (APA 7) — the "glue" must be on the
    # figure paragraph above, not on the caption itself. Applied at fixup time.
    fc = _get_or_add_paragraph_style(doc, "TME Figure Caption")
    fc.font.name = "Georgia"
    fc.font.size = Pt(10)
    fc.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    fc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fc.paragraph_format.space_before = Pt(8)
    fc.paragraph_format.space_after = Pt(18)

    tc = _get_or_add_paragraph_style(doc, "TME Table Caption")
    tc.font.name = "Georgia"
    tc.font.size = Pt(10)
    tc.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    tc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tc.paragraph_format.space_before = Pt(18)
    tc.paragraph_format.space_after = Pt(8)
    tc.paragraph_format.keep_with_next = True

    fn = _get_or_add_paragraph_style(doc, "TME Footnote")
    fn.font.name = "Georgia"
    fn.font.size = Pt(9)
    fn.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    fn.paragraph_format.line_spacing = 1.2
    fn.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ref = _get_or_add_paragraph_style(doc, "TME Reference")
    ref.font.name = "Georgia"
    ref.font.size = Pt(10.5)
    # Hanging indent: 1.5em at 10.5pt = 15.75pt
    ref.paragraph_format.left_indent = Pt(15.75)
    ref.paragraph_format.first_line_indent = Pt(-15.75)
    ref.paragraph_format.line_spacing = 1.45
    ref.paragraph_format.space_after = Pt(4)

    pq = _get_or_add_paragraph_style(doc, "TME Pullquote")
    pq.font.name = "Georgia"
    pq.font.size = Pt(15)
    pq.font.italic = True
    pq.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    pq.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pq.paragraph_format.line_spacing = 1.5
    pq.paragraph_format.space_before = Pt(22)
    pq.paragraph_format.space_after = Pt(22)
    # Horizontal rules above/below applied at element time — see oxml_helpers.

    bq = _get_or_add_paragraph_style(doc, "TME Block Quote")
    bq.font.name = "Georgia"
    bq.font.size = Pt(10.5)
    bq.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    bq.paragraph_format.left_indent = Pt(28)
    bq.paragraph_format.right_indent = Pt(28)
    bq.paragraph_format.line_spacing = 1.0
    bq.paragraph_format.space_before = Pt(8)
    bq.paragraph_format.space_after = Pt(8)

    # Word auto-creates "List Paragraph" when a list is pasted in. Pre-register
    # it so our body settings win over Word's defaults (which include 2.0 line
    # spacing in some source docs).
    lp = _get_or_add_paragraph_style(doc, "List Paragraph")
    lp.font.name = "Georgia"
    lp.font.size = Pt(11.5)
    lp.paragraph_format.line_spacing = 1.25
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after = Pt(4)
