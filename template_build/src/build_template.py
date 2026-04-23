"""Build TME_Template_2026.docx — the reusable template file."""
from pathlib import Path

from docx import Document
from docx.shared import Inches

from tme_template.cover_footer import add_cover_footer
from tme_template.cover_page import (
    AuthorEntry, CoverData, add_research_article_cover,
)
from tme_template.front_matter import (
    IssueInfo, StaffRoster,
    add_issue_cover_page, add_editorial_staff_page, add_formal_title_page,
)
from tme_template.headers_footers import set_running_footer, set_running_headers
from tme_template.masthead import MastheadData, add_masthead
from tme_template.oxml_helpers import add_section_break_next_page, add_continuous_section_break
from tme_template.page_setup import configure_page_setup, configure_zero_margins
from tme_template.styles import (
    register_body_style, register_heading_styles,
    register_remaining_styles, register_title_style,
)
from tme_template.tagline import add_tagline_strip


REPO_ROOT = Path(__file__).resolve().parents[2]
LOGO_H = str(REPO_ROOT / "assets" / "tme-logo.jpg")
LOGO_P = str(REPO_ROOT / "assets" / "tme-logo-portrait.jpg")
OUTPUT = REPO_ROOT / "TME_Template_2026.docx"


def build() -> Path:
    doc = Document()
    configure_page_setup(doc)
    register_body_style(doc)
    register_title_style(doc)
    register_heading_styles(doc)
    register_remaining_styles(doc)

    issue = IssueInfo(
        volume=34, number=1, year=2026, season="Spring",
        cover_artist="Keith Przybyla-Kuchek",
        portrait_logo_path=LOGO_P,
    )

    # --- Page 1: issue cover ---
    add_issue_cover_page(doc, issue)
    # No explicit page break: the full-page table fills the page on its own.

    # --- Page 2: formal title ---
    add_formal_title_page(doc, issue)

    # --- Page 3: editorial staff (masthead = zero-margin section, roster = normal) ---
    ed_masthead_section = add_section_break_next_page(doc)
    configure_zero_margins(ed_masthead_section)
    ed_masthead_section.page_width = Inches(8.5)
    ed_masthead_section.page_height = Inches(11)
    add_masthead(doc, MastheadData(
        article_type="EDITORIAL STAFF",
        volume=34, number=1, year=2026, pages=None,
        doi=None,
        issn_print="1062-9017",
        issn_online="2331-4451",
        logo_path=LOGO_H,
    ))
    add_tagline_strip(doc)
    ed_body_section = add_continuous_section_break(doc)
    ed_body_section.page_width = Inches(8.5)
    ed_body_section.page_height = Inches(11)
    ed_body_section.top_margin = Inches(0.3)
    ed_body_section.bottom_margin = Inches(0.5)
    ed_body_section.left_margin = Inches(0.5)
    ed_body_section.right_margin = Inches(0.5)
    ed_body_section.footer.is_linked_to_previous = False
    for p in list(ed_body_section.footer.paragraphs):
        p._p.getparent().remove(p._p)
    add_editorial_staff_page(doc, issue, StaffRoster(
        editors=["Jennifer Kleiman", "[Co-Editor Name]"],
        associate_editors=["[Associate Editor]", "[Associate Editor]"],
        advisor="[Advisor]",
        copy_editor="[Copy Editor]",
        mesa_officers={
            "President": "[President]",
            "Vice-President": "[VP]",
            "Secretary": "[Secretary]",
            "Treasurer": "[Treasurer]",
            "NCTM Representative": "[NCTM Rep]",
            "Colloquium Chair": "[Chair]",
        },
        mesa_term="2026-2027",
    ))

    # --- Page 4: sample research-article cover ---
    # Zero-margin section for full-bleed masthead + tagline + cover footer
    masthead_section = add_section_break_next_page(doc)
    configure_zero_margins(masthead_section)
    masthead_section.bottom_margin = Inches(0.5)   # give the footer room to render
    masthead_section.page_width = Inches(8.5)
    masthead_section.page_height = Inches(11)
    add_masthead(doc, MastheadData(
        article_type="RESEARCH ARTICLE",
        volume=34, number=1, year=2026, pages="1–24",
        doi="doi.org/10.xxxxx/tme.2026.34.1.01",
        issn_print="1062-9017", issn_online="2331-4451",
        logo_path=LOGO_H,
    ))
    add_tagline_strip(doc)
    add_cover_footer(masthead_section,
        citation="[Author, A. (YYYY). Title. The Mathematics Educator, V(N), pp–pp.]",
        license_text="CC BY 4.0",
        copyright_text="© 2026 The Authors",
    )
    cover_body_section = add_continuous_section_break(doc)
    cover_body_section.page_width = Inches(8.5)
    cover_body_section.page_height = Inches(11)
    cover_body_section.top_margin = Inches(0.3)
    cover_body_section.bottom_margin = Inches(0.3)
    cover_body_section.left_margin = Inches(0.5)
    cover_body_section.right_margin = Inches(0.5)
    cover_body_section.footer.is_linked_to_previous = False
    for p in list(cover_body_section.footer.paragraphs):
        p._p.getparent().remove(p._p)
    add_research_article_cover(doc, CoverData(
        title="[Article Title Goes Here]",
        authors=[
            AuthorEntry(name="[Author One]", affiliation_num=1,
                        role="[Role, Institution]",
                        bio="[Bio paragraph.]",
                        headshot_path=None,
                        corresponding=True,
                        email="[email@uga.edu]"),
        ],
        affiliations=["[Affiliation]"],
        dates={"Received": "[Date]", "Revised": "[Date]",
               "Accepted": "[Date]", "Published": "[Date]"},
        abstract="[Abstract paragraph goes here.]",
        keywords=["[keyword one]", "[keyword two]"],
    ))

    # --- Final section: body pages with running headers + footer ---
    # A NEW_PAGE section break begins the body pages on a fresh page.
    body_section = add_section_break_next_page(doc)

    set_running_headers(doc,
        author_cite="[Short Author Cite]",
        short_title="[Short Title]",
        section=body_section)
    set_running_footer(doc, section=body_section)

    # --- Sample body page ---
    doc.add_paragraph("[Body Heading]", style="TME H1")
    doc.add_paragraph("[First body paragraph of the article.]", style="TME Body")
    doc.add_paragraph("[Second paragraph.]", style="TME Body")

    doc.save(str(OUTPUT))
    return OUTPUT


if __name__ == "__main__":
    out = build()
    print(f"Wrote {out}")
