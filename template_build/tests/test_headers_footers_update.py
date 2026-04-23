"""Assert the body running footer is just a centered page number."""
from docx import Document

from tme_template.headers_footers import set_running_footer


def _footer_text(section):
    parts = []
    for p in section.footer.paragraphs:
        parts.append(p.text)
    return " | ".join(parts)


def _footer_has_page_field(section):
    """Search the footer XML for a PAGE field instruction."""
    xml = section.footer._element.xml
    return 'PAGE' in xml and 'w:fldChar' in xml


def test_footer_has_no_copyright_text():
    doc = Document()
    section = doc.sections[0]
    set_running_footer(doc, section=section)
    txt = _footer_text(section)
    assert "©" not in txt
    assert "CC BY" not in txt
    assert "Authors" not in txt


def test_footer_contains_page_field():
    doc = Document()
    section = doc.sections[0]
    set_running_footer(doc, section=section)
    assert _footer_has_page_field(section)


def test_footer_paragraph_is_centered():
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    section = doc.sections[0]
    set_running_footer(doc, section=section)
    assert section.footer.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
