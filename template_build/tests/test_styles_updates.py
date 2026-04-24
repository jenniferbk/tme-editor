"""Assert the post-Moore-proof style updates land correctly."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from tme_template.styles import (
    register_title_style,
    register_heading_styles,
    register_remaining_styles,
)


def test_title_has_24pt_space_before():
    doc = Document()
    register_title_style(doc)
    style = doc.styles["TME Title"]
    assert style.paragraph_format.space_before == Pt(24)


def test_h3_is_bold_not_italic():
    doc = Document()
    register_heading_styles(doc)
    h3 = doc.styles["TME H3"]
    assert h3.font.bold is True
    assert h3.font.italic is False


def test_h1_and_h2_italicization_unchanged():
    doc = Document()
    register_heading_styles(doc)
    h1 = doc.styles["TME H1"]
    h2 = doc.styles["TME H2"]
    assert h1.font.bold is True and h1.font.italic is not True
    assert h2.font.bold is True and h2.font.italic is True


def test_figure_caption_is_left_aligned_and_sticky():
    doc = Document()
    register_remaining_styles(doc)
    fc = doc.styles["TME Figure Caption"]
    assert fc.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert fc.paragraph_format.keep_with_next is True


def test_table_caption_is_left_aligned_and_sticky():
    doc = Document()
    register_remaining_styles(doc)
    tc = doc.styles["TME Table Caption"]
    assert tc.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert tc.paragraph_format.keep_with_next is True
