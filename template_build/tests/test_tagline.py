from docx import Document
from docx.oxml.ns import qn
from tme_template.tagline import add_tagline_strip


def test_tagline_strip_creates_single_cell_table():
    doc = Document()
    add_tagline_strip(doc)
    table = doc.tables[0]
    assert len(table.rows) == 1
    assert len(table.columns) == 1


def test_tagline_cell_has_light_panel_gray_shading():
    doc = Document()
    add_tagline_strip(doc)
    cell = doc.tables[0].cell(0, 0)
    shd = cell._tc.tcPr.find(qn("w:shd"))
    assert shd.get(qn("w:fill")) == "F5F5F5"


def test_tagline_cell_contains_tagline_text_and_meta_line():
    doc = Document()
    add_tagline_strip(doc)
    cell = doc.tables[0].cell(0, 0)
    text = "\n".join(p.text for p in cell.paragraphs)
    assert "Cultivating scholarly discourse" in text
    assert "1990" in text
    assert "Mathematics Education Student Association" in text
    assert "University of Georgia" in text
    assert "Peer Reviewed" in text
    assert "Open Access" in text


def test_tagline_contains_red_diamond_glyphs():
    doc = Document()
    add_tagline_strip(doc)
    cell = doc.tables[0].cell(0, 0)
    text = "\n".join(p.text for p in cell.paragraphs)
    assert "◆" in text
