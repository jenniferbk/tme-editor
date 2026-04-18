from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH

from tme_template.headers_footers import (
    set_running_headers,
    set_running_footer,
)
from tme_template.page_setup import configure_page_setup


def test_verso_header_is_left_aligned_with_author_cite():
    doc = Document()
    configure_page_setup(doc)
    set_running_headers(doc, author_cite="Moore, Yasuda, & Wong",
                        short_title="Integration by Substitution")
    section = doc.sections[0]
    verso = section.even_page_header
    assert "Moore, Yasuda, & Wong" in verso.paragraphs[0].text
    assert verso.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT


def test_recto_header_is_right_aligned_with_short_title():
    doc = Document()
    configure_page_setup(doc)
    set_running_headers(doc, author_cite="Moore, Yasuda, & Wong",
                        short_title="Integration by Substitution")
    section = doc.sections[0]
    recto = section.header  # default (odd) header
    assert "Integration by Substitution" in recto.paragraphs[0].text
    assert recto.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT


def test_set_running_headers_on_section2_leaves_section1_empty():
    """Setting headers on a specific section must not bleed into the prior section.

    A paragraph must exist before calling add_section so that python-docx places
    section1's sectPr inline (in a paragraph's pPr), giving each section its own
    distinct sectPr element.  Without content, both sections share the body's
    trailing sectPr and the isolation cannot be tested meaningfully.
    """
    doc = Document()
    configure_page_setup(doc)
    doc.add_paragraph("Section 1 placeholder")  # ensures section1 gets its own sectPr
    doc.add_section(WD_SECTION.NEW_PAGE)

    section1 = doc.sections[0]
    section2 = doc.sections[1]

    set_running_headers(doc, author_cite="Moore & Yasuda",
                        short_title="[Short Title]", section=section2)

    # section2 should have content in both odd and even headers
    assert "[Short Title]" in section2.header.paragraphs[0].text
    assert "Moore & Yasuda" in section2.even_page_header.paragraphs[0].text

    # section1's odd and even headers should still be empty (no text)
    s1_odd_text = "".join(p.text for p in section1.header.paragraphs)
    s1_even_text = "".join(p.text for p in section1.even_page_header.paragraphs)
    assert s1_odd_text == "", f"section1 odd header unexpectedly has text: {s1_odd_text!r}"
    assert s1_even_text == "", f"section1 even header unexpectedly has text: {s1_even_text!r}"


def test_running_footer_has_page_number_and_license():
    doc = Document()
    configure_page_setup(doc)
    set_running_footer(doc, copyright_line="© 2026 The Authors  ·  CC BY 4.0")
    section = doc.sections[0]
    # Odd footer: license left, page number right
    odd = section.footer
    assert "© 2026" in odd.paragraphs[0].text
    # {PAGE} field appears as empty before rendering, but the OXML should have it
    from docx.oxml.ns import qn
    fld = odd.paragraphs[0]._p.xpath(".//w:instrText")
    assert any("PAGE" in (f.text or "") for f in fld)
