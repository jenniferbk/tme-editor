from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

from tme_template.styles import (
    register_body_style,
    register_title_style,
)


def test_body_style_registered_with_correct_properties():
    doc = Document()
    register_body_style(doc)
    style = doc.styles["TME Body"]
    assert style.type == WD_STYLE_TYPE.PARAGRAPH
    assert style.font.name == "Georgia"
    assert style.font.size == Pt(11.5)
    # Line height 1.65 → line_spacing attribute
    assert style.paragraph_format.line_spacing == 1.65
    # First-line indent 1.5em → approx 0.25" for 11.5pt body
    assert style.paragraph_format.first_line_indent is not None


def test_title_style_registered():
    doc = Document()
    register_title_style(doc)
    style = doc.styles["TME Title"]
    assert style.font.name == "Georgia"
    assert style.font.size == Pt(18)
    assert style.font.bold is True


from tme_template.styles import register_heading_styles


def test_h1_style():
    doc = Document()
    register_heading_styles(doc)
    h1 = doc.styles["TME H1"]
    assert h1.font.name == "Georgia"
    assert h1.font.size == Pt(16)
    assert h1.font.bold is True


def test_h2_style():
    doc = Document()
    register_heading_styles(doc)
    h2 = doc.styles["TME H2"]
    assert h2.font.size == Pt(13)
    assert h2.font.bold is True
    assert h2.font.italic is True


def test_h3_style():
    doc = Document()
    register_heading_styles(doc)
    h3 = doc.styles["TME H3"]
    assert h3.font.size == Pt(11.5)
    assert h3.font.bold is False
    assert h3.font.italic is True


from tme_template.styles import register_remaining_styles


def test_figure_caption_style():
    doc = Document()
    register_remaining_styles(doc)
    s = doc.styles["TME Figure Caption"]
    assert s.font.name == "Arial"
    assert s.font.size == Pt(10)


def test_table_caption_style():
    doc = Document()
    register_remaining_styles(doc)
    s = doc.styles["TME Table Caption"]
    assert s.font.name == "Arial"
    assert s.font.size == Pt(10)


def test_footnote_text_style():
    doc = Document()
    register_remaining_styles(doc)
    s = doc.styles["TME Footnote"]
    assert s.font.name == "Georgia"
    assert s.font.size == Pt(9)


def test_reference_entry_style():
    doc = Document()
    register_remaining_styles(doc)
    s = doc.styles["TME Reference"]
    assert s.font.name == "Georgia"
    assert s.font.size == Pt(10.5)
    # Hanging indent: 1.5em ≈ 15.75pt at 10.5pt
    assert s.paragraph_format.left_indent is not None
    assert s.paragraph_format.first_line_indent is not None


def test_pullquote_style():
    doc = Document()
    register_remaining_styles(doc)
    s = doc.styles["TME Pullquote"]
    assert s.font.name == "Georgia"
    assert s.font.size == Pt(15)
    assert s.font.italic is True


def test_blockquote_style():
    doc = Document()
    register_remaining_styles(doc)
    s = doc.styles["TME Block Quote"]
    assert s.font.name == "Georgia"
    assert s.font.size == Pt(10.5)
