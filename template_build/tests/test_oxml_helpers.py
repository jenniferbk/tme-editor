from docx import Document
from docx.oxml.ns import qn

from tme_template.oxml_helpers import (
    set_cell_shading,
    remove_cell_borders,
    set_different_odd_even_pages,
    set_different_first_page,
)


def test_set_cell_shading_adds_shd_element():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "BA0C2F")
    shd = cell._tc.tcPr.find(qn("w:shd"))
    assert shd is not None
    assert shd.get(qn("w:fill")) == "BA0C2F"
    assert shd.get(qn("w:val")) == "clear"


def test_remove_cell_borders_sets_all_sides_to_nil():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    remove_cell_borders(cell)
    tcBorders = cell._tc.tcPr.find(qn("w:tcBorders"))
    assert tcBorders is not None
    for side in ("top", "left", "bottom", "right"):
        side_el = tcBorders.find(qn(f"w:{side}"))
        assert side_el is not None
        assert side_el.get(qn("w:val")) == "nil"


def test_set_different_odd_even_pages_adds_setting():
    doc = Document()
    set_different_odd_even_pages(doc)
    settings = doc.settings.element
    evenAndOddHeaders = settings.find(qn("w:evenAndOddHeaders"))
    assert evenAndOddHeaders is not None


def test_set_different_first_page_on_section():
    doc = Document()
    section = doc.sections[0]
    set_different_first_page(section, True)
    assert section.different_first_page_header_footer is True


from docx import Document
from tme_template.oxml_helpers import (
    apply_red_left_rule,
    apply_bottom_rule,
    apply_top_rule,
    apply_pullquote_rules,
)
from docx.oxml.ns import qn


def test_apply_red_left_rule_adds_left_border():
    doc = Document()
    p = doc.add_paragraph("Hello")
    apply_red_left_rule(p, hex_color="BA0C2F", width_pt=3)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    assert pBdr is not None
    left = pBdr.find(qn("w:left"))
    assert left is not None
    assert left.get(qn("w:color")) == "BA0C2F"


def test_apply_bottom_rule_adds_bottom_border():
    doc = Document()
    p = doc.add_paragraph("Rule test")
    apply_bottom_rule(p, hex_color="CCCCCC", width_pt=1)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    assert pBdr is not None
    bottom = pBdr.find(qn("w:bottom"))
    assert bottom is not None
    assert bottom.get(qn("w:color")) == "CCCCCC"
    assert bottom.get(qn("w:sz")) == "8"  # 1pt * 8 = 8 eighths


def test_apply_top_rule_adds_top_border():
    doc = Document()
    p = doc.add_paragraph("Top rule test")
    apply_top_rule(p, hex_color="EEEEEE", width_pt=1)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    assert pBdr is not None
    top = pBdr.find(qn("w:top"))
    assert top is not None
    assert top.get(qn("w:color")) == "EEEEEE"
    assert top.get(qn("w:sz")) == "8"  # 1pt * 8 = 8 eighths


def test_apply_pullquote_rules_adds_top_and_bottom_borders():
    doc = Document()
    p = doc.add_paragraph("Quote")
    apply_pullquote_rules(p, top_hex="BA0C2F", bottom_hex="BA0C2F")
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    top = pBdr.find(qn("w:top"))
    bottom = pBdr.find(qn("w:bottom"))
    assert top is not None
    assert bottom is not None
    assert top.get(qn("w:color")) == "BA0C2F"
    assert bottom.get(qn("w:color")) == "BA0C2F"
