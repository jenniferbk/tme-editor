from dataclasses import replace
from docx import Document
from tme_template.styles import register_body_style, register_title_style, register_remaining_styles
from tme_template.cover_page import add_research_article_cover, CoverData, AuthorEntry


def _authors():
    return [
        AuthorEntry(name="Kevin C. Moore", affiliation_num=1,
                    role="Professor, UGA",
                    bio="Studies quantitative reasoning in calculus.",
                    headshot_path=None,
                    corresponding=True,
                    email="kvcmoore@uga.edu"),
        AuthorEntry(name="Sohei Yasuda", affiliation_num=1,
                    role="Doctoral student, UGA",
                    bio="Research on multivariable calculus.",
                    headshot_path=None),
        AuthorEntry(name="Webster Wong", affiliation_num=1,
                    role="Doctoral student, UGA",
                    bio="Undergraduate math curriculum research.",
                    headshot_path=None),
    ]


def _cover() -> CoverData:
    return CoverData(
        title="Integration by Substitution: An Emergent Quantitative Reasoning Approach",
        authors=_authors(),
        affiliations=["Department of Mathematics and Science Education, University of Georgia"],
        dates={"Received": "Dec 3, 2024", "Revised": "Jan 22, 2025",
               "Accepted": "Apr 3, 2025", "Published": "Apr 2026"},
        abstract="In this paper, we present a conceptual analysis...",
        keywords=["Conceptual Analysis", "Integration by Substitution"],
    )


def _all_text(doc) -> str:
    """Collect text from all paragraphs and all table cell paragraphs."""
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for cell in tbl._cells:
            for p in cell.paragraphs:
                parts.append(p.text)
    return "\n".join(parts)


def test_cover_includes_title_authors_affiliations_dates_abstract_keywords():
    doc = Document()
    register_body_style(doc); register_title_style(doc); register_remaining_styles(doc)
    add_research_article_cover(doc, _cover())
    body_text = _all_text(doc)
    assert "Integration by Substitution" in body_text
    assert "Kevin C. Moore" in body_text
    assert "University of Georgia" in body_text
    assert "RECEIVED" in body_text
    assert "Dec 3, 2024" in body_text
    assert "conceptual analysis" in body_text.lower()
    assert "Keywords" in body_text


def test_corresponding_author_has_dagger_marker_and_email_line():
    doc = Document()
    register_body_style(doc); register_title_style(doc); register_remaining_styles(doc)
    add_research_article_cover(doc, _cover())
    text = _all_text(doc)
    assert "†" in text
    assert "kvcmoore@uga.edu" in text


def test_author_block_label_is_present():
    doc = Document()
    register_body_style(doc); register_title_style(doc); register_remaining_styles(doc)
    add_research_article_cover(doc, _cover())
    text = _all_text(doc)
    assert "ABOUT THE AUTHORS" in text


def test_author_block_is_table_with_one_cell_per_author():
    """Author block must be rendered as a 3-column table row, not stacked paragraphs."""
    doc = Document()
    register_body_style(doc); register_title_style(doc); register_remaining_styles(doc)
    add_research_article_cover(doc, _cover())
    # Find the author table (should be the last table in the document)
    assert len(doc.tables) >= 1
    author_table = doc.tables[-1]
    assert len(author_table.columns) == 3  # one column per author
    assert len(author_table.rows) == 1
    # Each author name should appear inside the table cells
    cell_text = "\n".join(
        p.text for cell in author_table._cells for p in cell.paragraphs
    )
    assert "Kevin C. Moore" in cell_text
    assert "Sohei Yasuda" in cell_text
    assert "Webster Wong" in cell_text
