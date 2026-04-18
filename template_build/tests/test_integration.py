"""End-to-end integration test: generate template and verify structural properties."""
import subprocess
from pathlib import Path

import pytest
from docx import Document


# Derive paths relative to this test file:
# tests/test_integration.py → tests/ → template_build/ → TME/
_TEST_DIR = Path(__file__).resolve().parent
_TEMPLATE_BUILD_DIR = _TEST_DIR.parent
_TME_DIR = _TEMPLATE_BUILD_DIR.parent

OUTPUT = _TME_DIR / "TME_Template_2026.docx"


@pytest.fixture(scope="module", autouse=True)
def _build_template():
    """Run the build script once per test module."""
    subprocess.run(
        ["python3", "src/build_template.py"],
        cwd=str(_TEMPLATE_BUILD_DIR),
        check=True,
    )
    yield
    # Don't delete — user inspects the output.


def test_template_has_all_named_styles():
    doc = Document(str(OUTPUT))
    names = {s.name for s in doc.styles}
    for required in [
        "TME Body", "TME Title",
        "TME H1", "TME H2", "TME H3",
        "TME Figure Caption", "TME Table Caption",
        "TME Footnote", "TME Reference",
        "TME Pullquote", "TME Block Quote",
    ]:
        assert required in names, f"Missing style: {required}"


def test_template_has_correct_page_size_and_margins():
    from docx.shared import Inches
    doc = Document(str(OUTPUT))
    s = doc.sections[0]
    assert abs(s.page_width - Inches(8.5)) < 1000
    assert abs(s.page_height - Inches(11)) < 1000
    # Top/bottom are tightened to 0.3"; left/right remain 0.5"
    assert abs(s.top_margin - Inches(0.3)) < 1000
    assert abs(s.bottom_margin - Inches(0.3)) < 1000
    assert abs(s.left_margin - Inches(0.5)) < 1000
    assert abs(s.right_margin - Inches(0.5)) < 1000


def test_template_has_odd_even_pages_enabled():
    from docx.oxml.ns import qn
    doc = Document(str(OUTPUT))
    assert doc.settings.element.find(qn("w:evenAndOddHeaders")) is not None


def test_template_masthead_has_uga_red_and_black():
    from docx.oxml.ns import qn
    doc = Document(str(OUTPUT))
    # Multiple mastheads — find one with UGA_RED shading and BLACK shading
    found_red = False
    found_black = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                shd = cell._tc.tcPr.find(qn("w:shd")) if cell._tc.tcPr is not None else None
                if shd is None:
                    continue
                fill = shd.get(qn("w:fill"))
                if fill == "BA0C2F":
                    found_red = True
                if fill == "000000":
                    found_black = True
    assert found_red, "No cell with UGA red shading"
    assert found_black, "No cell with black shading"


def test_template_contains_all_five_article_type_mentions_or_one_and_scaffolding():
    """We don't require all five article types in a *single* template file —
    the template is a starting point. But it should at least contain the
    research-article masthead, the editorial-staff masthead, and the tagline strip."""
    doc = Document(str(OUTPUT))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for cell in table._cells:
            for p in cell.paragraphs:
                all_text += "\n" + p.text
    assert "RESEARCH ARTICLE" in all_text
    assert "EDITORIAL STAFF" in all_text
    assert "Cultivating scholarly discourse" in all_text


def test_template_opens_without_xml_errors():
    """Opening and re-saving should round-trip cleanly."""
    doc = Document(str(OUTPUT))
    # Round-trip save to a tmp path
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        doc.save(tf.name)
        # Re-open to verify
        doc2 = Document(tf.name)
        assert len(doc2.paragraphs) > 0
