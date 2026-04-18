from docx import Document
from tme_template.front_matter import (
    add_issue_cover_page,
    add_editorial_staff_page,
    add_formal_title_page,
    IssueInfo, StaffRoster,
)


def _issue():
    return IssueInfo(volume=34, number=1, year=2026, season="Spring",
                     cover_artist="Jane Doe", portrait_logo_path="/tmp/does-not-exist.jpg")


def _all_text(doc) -> str:
    """Collect text from paragraphs and table cell paragraphs."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return "\n".join(parts)


def test_issue_cover_page_has_volume_season_and_artist_credit():
    doc = Document()
    add_issue_cover_page(doc, _issue())
    text = _all_text(doc)
    assert "Volume 34" in text or "VOLUME 34" in text
    assert "Spring 2026" in text
    assert "Jane Doe" in text


def test_editorial_staff_page_lists_roles():
    roster = StaffRoster(
        editors=["Jennifer Kleiman", "Co-Editor"],
        associate_editors=["Associate A", "Associate B"],
        advisor="Advisor Name",
        copy_editor="Copy Editor Name",
        mesa_officers={
            "President": "President Name",
            "Vice-President": "VP Name",
            "Secretary": "Sec Name",
            "Treasurer": "Treasurer Name",
        },
        mesa_term="2026-2027",
    )
    doc = Document()
    add_editorial_staff_page(doc, _issue(), roster)
    text = _all_text(doc)
    assert "EDITORIAL BOARD" in text
    assert "Jennifer Kleiman" in text
    assert "Associate A" in text
    assert "MESA OFFICERS" in text
    assert "President Name" in text


def test_formal_title_page_has_wordmark_and_issue_info():
    doc = Document()
    add_formal_title_page(doc, _issue())
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "THE MATHEMATICS EDUCATOR" in text
    assert "Mathematics Education Student Association" in text
    assert "University of Georgia" in text
    assert "VOLUME 34" in text
    assert "Spring 2026" in text
