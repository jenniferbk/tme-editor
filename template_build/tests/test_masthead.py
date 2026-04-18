from docx import Document
from docx.oxml.ns import qn

from tme_template.masthead import add_masthead, MastheadData


def _sample() -> MastheadData:
    return MastheadData(
        article_type="RESEARCH ARTICLE",
        volume=34, number=1, year=2026,
        pages="1–24",
        doi="doi.org/10.xxxxx/tme.2026.34.1.01",
        issn_print="1062-9017",
        issn_online="2331-4451",
        logo_path="/Users/jenniferkleiman/Documents/TME/.superpowers/brainstorm/5088-1776377525/content/tme-logo.jpg",
    )


def test_masthead_creates_2col_table():
    doc = Document()
    add_masthead(doc, _sample())
    table = doc.tables[0]
    assert len(table.rows) == 1
    assert len(table.columns) == 2


def test_right_cell_has_uga_red_shading():
    doc = Document()
    add_masthead(doc, _sample())
    right_cell = doc.tables[0].cell(0, 1)
    shd = right_cell._tc.tcPr.find(qn("w:shd"))
    assert shd.get(qn("w:fill")) == "BA0C2F"


def test_left_cell_has_black_shading():
    doc = Document()
    add_masthead(doc, _sample())
    left_cell = doc.tables[0].cell(0, 0)
    shd = left_cell._tc.tcPr.find(qn("w:shd"))
    assert shd.get(qn("w:fill")) == "000000"


def test_right_cell_contains_article_type_label():
    doc = Document()
    add_masthead(doc, _sample())
    right_cell = doc.tables[0].cell(0, 1)
    text = "\n".join(p.text for p in right_cell.paragraphs)
    assert "RESEARCH ARTICLE" in text
    assert "Vol. 34, No. 1 (2026), pp. 1–24" in text
    assert "doi.org/10.xxxxx/tme.2026.34.1.01" in text
    assert "1062-9017" in text
    assert "2331-4451" in text


def test_missing_doi_leaves_slot_blank_but_present():
    """For article types with no DOI (Editorial Staff, TOC) the slot is blank,
    preserving masthead height."""
    doc = Document()
    data = MastheadData(
        article_type="EDITORIAL STAFF",
        volume=34, number=1, year=2026, pages=None,
        doi=None, issn_print="1062-9017", issn_online="2331-4451",
        logo_path="/tmp/does-not-exist.jpg",
    )
    add_masthead(doc, data)
    right_cell = doc.tables[0].cell(0, 1)
    # Should have 4 paragraphs (type, vol line, empty doi slot, issn)
    assert len(right_cell.paragraphs) == 4
