from docx import Document
from docx.oxml.ns import qn
from tme_template.cover_footer import add_cover_footer


def _make_section():
    """Create a fresh document and return its first section."""
    doc = Document()
    return doc, doc.sections[0]


def test_cover_footer_creates_2col_table():
    doc, section = _make_section()
    add_cover_footer(section, citation="Moore, K. C. (2026). Title. TME 34(1), 1–24.",
                     license_text="CC BY 4.0", copyright_text="© 2026 The Authors")
    tables = list(section.footer.tables)
    assert len(tables) == 1
    table = tables[0]
    assert len(table.rows) == 1
    assert len(table.columns) == 2


def test_cover_footer_left_cell_has_how_to_cite_label():
    doc, section = _make_section()
    add_cover_footer(section, citation="Moore...", license_text="CC BY 4.0",
                     copyright_text="© 2026 The Authors")
    table = list(section.footer.tables)[0]
    left = table.cell(0, 0)
    text = "\n".join(p.text for p in left.paragraphs)
    assert "HOW TO CITE" in text
    assert "Moore" in text


def test_cover_footer_right_cell_has_license_and_copyright():
    doc, section = _make_section()
    add_cover_footer(section, citation="...", license_text="CC BY 4.0",
                     copyright_text="© 2026 The Authors")
    table = list(section.footer.tables)[0]
    right = table.cell(0, 1)
    text = "\n".join(p.text for p in right.paragraphs)
    assert "CC BY 4.0" in text
    assert "2026 The Authors" in text


def test_cover_footer_has_cream_background():
    doc, section = _make_section()
    add_cover_footer(section, citation="...", license_text="CC BY 4.0",
                     copyright_text="© 2026 The Authors")
    table = list(section.footer.tables)[0]
    for i in (0, 1):
        cell = table.cell(0, i)
        shd = cell._tc.tcPr.find(qn("w:shd"))
        assert shd.get(qn("w:fill")) == "FAFAF7"
