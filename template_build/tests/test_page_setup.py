from docx import Document
from docx.shared import Inches
from tme_template.page_setup import configure_page_setup


def test_page_is_letter_size():
    doc = Document()
    configure_page_setup(doc)
    s = doc.sections[0]
    assert abs(s.page_width - Inches(8.5)) < 1000
    assert abs(s.page_height - Inches(11)) < 1000


def test_margins():
    doc = Document()
    configure_page_setup(doc)
    s = doc.sections[0]
    assert abs(s.top_margin - Inches(0.3)) < 1000
    assert abs(s.bottom_margin - Inches(0.3)) < 1000
    assert abs(s.left_margin - Inches(0.5)) < 1000
    assert abs(s.right_margin - Inches(0.5)) < 1000


def test_odd_even_headers_enabled():
    from docx.oxml.ns import qn
    doc = Document()
    configure_page_setup(doc)
    evenOdd = doc.settings.element.find(qn("w:evenAndOddHeaders"))
    assert evenOdd is not None
