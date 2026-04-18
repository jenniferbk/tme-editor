"""Apply a set of post-paste, post-styling fixes to TME_Moore_2026_starter.docx.

Run after apply_tme_styles.py. Backs up before mutating.

Fixes:
  1. Body line-spacing 1.65 -> 1.25
  2. Figure/Table Caption font Arial -> Georgia
  3. TME Footnote line-spacing 1.5 -> 1.2 (Georgia already)
  4. TME Block Quote line_spacing 1.0, space_before/after 8pt
  5. Remap body paragraphs with left_indent > 20pt -> TME Block Quote
  6. Center content tables, prevent mid-row splits, repeat first row as header
  7. Force Georgia on every run inside footnotes.xml
  8. Strip direct run-level font size/bold/rFonts on TME Reference paragraphs
  9. Rewrite masthead+tagline table grid to eliminate phantom edge columns
"""
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


TME = Path(__file__).resolve().parents[2]
STARTER = TME / "TME_Moore_2026_starter.docx"
BACKUP = STARTER.with_name("TME_Moore_2026_starter.pre-fixup.docx")


# ---------- style definition updates ----------

def update_styles(doc) -> None:
    styles = doc.styles
    style_names = {s.name for s in styles}

    body = styles["TME Body"]
    body.paragraph_format.line_spacing = 1.25

    # Heading styles: keep heading with next paragraph so it doesn't orphan
    # at the bottom of a page.
    for hname in ("TME H1", "TME H2", "TME H3"):
        if hname in style_names:
            h = styles[hname]
            h.paragraph_format.keep_with_next = True
            h.paragraph_format.keep_together = True

    # List Paragraph (Word default): pasted Moore lists inherit 2.0 line spacing
    # via direct paragraph formatting. Fix the style definition here; we'll
    # also strip direct line_spacing on each instance below.
    if "List Paragraph" in style_names:
        lp = styles["List Paragraph"]
        lp.font.name = "Georgia"
        lp.font.size = Pt(11.5)
        lp.paragraph_format.line_spacing = 1.25
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after = Pt(4)

    fc = styles["TME Figure Caption"]
    fc.font.name = "Georgia"

    tc = styles["TME Table Caption"]
    tc.font.name = "Georgia"

    fn = styles["TME Footnote"]
    fn.font.name = "Georgia"
    fn.paragraph_format.line_spacing = 1.2

    # Also update the Word-builtin footnote styles (inherited from the Moore
    # source). Footnote runs without explicit rPr fall back to these.
    # Name casing varies ("Footnote Text" vs "footnote text") — match case-insensitively.
    def _find_style(*candidates):
        cl = {n.lower(): n for n in style_names}
        for c in candidates:
            actual = cl.get(c.lower())
            if actual is not None:
                return styles[actual]
        return None

    ft = _find_style("Footnote Text", "footnote text")
    if ft is not None:
        ft.font.name = "Georgia"
        ft.font.size = Pt(9)
        if ft.paragraph_format is not None:
            ft.paragraph_format.line_spacing = 1.2

    for sname in ("Footnote Reference", "footnote reference", "Footnote Text Char"):
        s = _find_style(sname)
        if s is not None:
            s.font.name = "Georgia"

    # TME Block Quote may or may not exist in the doc depending on template version.
    if "TME Block Quote" in style_names:
        bq = styles["TME Block Quote"]
    else:
        from docx.enum.style import WD_STYLE_TYPE
        bq = styles.add_style("TME Block Quote", WD_STYLE_TYPE.PARAGRAPH)
    bq.font.name = "Georgia"
    bq.font.size = Pt(10.5)
    bq.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    bq.paragraph_format.left_indent = Pt(28)
    bq.paragraph_format.right_indent = Pt(28)
    bq.paragraph_format.line_spacing = 1.0
    bq.paragraph_format.space_before = Pt(8)
    bq.paragraph_format.space_after = Pt(8)


# ---------- block-quote remapping ----------

def remap_block_quotes(doc) -> int:
    """Body paragraphs pasted with a left indent > 20pt are block quotes."""
    bq_style = doc.styles["TME Block Quote"]
    n = 0
    for p in doc.paragraphs:
        if p.style.name != "TME Body":
            continue
        li = p.paragraph_format.left_indent
        if li is None:
            continue
        if li.pt > 20:
            p.style = bq_style
            # Clear the direct left_indent so the style's indent takes over.
            p.paragraph_format.left_indent = None
            p.paragraph_format.first_line_indent = None
            n += 1
    return n


def fix_caption_classifications(doc) -> dict:
    """Some captions came in as TME Body or as the wrong caption kind. Reclassify
    by leading word ('Figure' -> TME Figure Caption, 'Table' -> TME Table Caption).
    Only targets paragraphs with short length (captions), to avoid eating body
    text that happens to start with 'Figure'.
    """
    fc = doc.styles["TME Figure Caption"]
    tc = doc.styles["TME Table Caption"]
    stats = {"fig_from_body": 0, "tab_from_body": 0, "fig_fix": 0, "tab_fix": 0}
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t or len(t) > 400:
            continue
        starts_fig = t.startswith("Figure")
        starts_tab = t.startswith("Table")
        if not (starts_fig or starts_tab):
            continue
        sn = p.style.name
        if starts_fig and sn == "TME Body":
            p.style = fc
            stats["fig_from_body"] += 1
        elif starts_tab and sn == "TME Body":
            p.style = tc
            stats["tab_from_body"] += 1
        elif starts_fig and sn == "TME Table Caption":
            p.style = fc
            stats["fig_fix"] += 1
        elif starts_tab and sn == "TME Figure Caption":
            p.style = tc
            stats["tab_fix"] += 1
    return stats


def clear_list_direct_spacing(doc) -> int:
    """Moore's numbered lists were pasted with direct line_spacing=2.0. Clear
    that so the List Paragraph style's 1.25 takes effect."""
    n = 0
    for p in doc.paragraphs:
        if p.style.name != "List Paragraph":
            continue
        pf = p.paragraph_format
        changed = False
        if pf.line_spacing is not None:
            pf.line_spacing = None
            changed = True
        if pf.space_before is not None and pf.space_before.pt > 0:
            pf.space_before = None
            changed = True
        if changed:
            n += 1
    return n


# ---------- table fixes ----------

def _set_trPr_flag(tr, tag_name: str) -> None:
    """Ensure <w:{tag_name}/> exists inside trPr."""
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    existing = trPr.find(qn(f"w:{tag_name}"))
    if existing is None:
        el = OxmlElement(f"w:{tag_name}")
        trPr.append(el)


def fix_content_tables(doc, skip_indices=(0, 1)) -> int:
    """Center content tables, prevent row splits mid-page, repeat header row.

    Tables at skip_indices are masthead (0) and author-card row (1).
    """
    n = 0
    for i, table in enumerate(doc.tables):
        if i in skip_indices:
            continue
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Ensure tblInd is 0 so centering actually happens (override any stale indent)
        tblPr = table._tbl.find(qn("w:tblPr"))
        if tblPr is not None:
            tblInd = tblPr.find(qn("w:tblInd"))
            if tblInd is not None:
                tblInd.set(qn("w:w"), "0")
            # Explicit jc=center inside tblPr (belt and suspenders)
            jc = tblPr.find(qn("w:jc"))
            if jc is None:
                jc = OxmlElement("w:jc")
                tblPr.append(jc)
            jc.set(qn("w:val"), "center")

        rows = table._tbl.findall(qn("w:tr"))
        for tr in rows:
            _set_trPr_flag(tr, "cantSplit")
        # Repeat first row as header on subsequent pages
        if rows:
            _set_trPr_flag(rows[0], "tblHeader")
        n += 1
    return n


# ---------- masthead grid fix ----------

def fix_masthead_grid(doc) -> bool:
    """Rewrite the merged masthead+tagline table grid to eliminate phantom
    edge columns. Assumes table index 0 is that table.
    """
    if not doc.tables:
        return False
    table = doc.tables[0]
    tbl = table._tbl

    # 8.5" @ 38/62 split, plus ~0.06" bleed on the right so Word's
    # zero-margin rendering reaches the page edge in Compatibility Mode.
    BLEED = 90  # twips, ~0.063"
    COL0, COL1 = 4651, 7589 + BLEED
    TOTAL = COL0 + COL1

    # Replace tblGrid with 2 cols
    existing_grid = tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    tblGrid = OxmlElement("w:tblGrid")
    for w in (COL0, COL1):
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        tblGrid.append(col)
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl.insert(0, tblGrid)

    # Strip gridBefore/gridAfter/wBefore/wAfter and cantSplit defaults from every row
    rows = tbl.findall(qn("w:tr"))
    for tr in rows:
        trPr = tr.find(qn("w:trPr"))
        if trPr is not None:
            for tag in ("w:gridBefore", "w:gridAfter", "w:wBefore", "w:wAfter"):
                for el in trPr.findall(qn(tag)):
                    trPr.remove(el)

    # Fix up each cell's tcW and gridSpan.
    # Row 0 (masthead): 2 cells -> tcW = [COL0, COL1], no gridSpan.
    # Row 1 (tagline): 1 cell -> tcW = TOTAL, gridSpan = 2.
    if len(rows) >= 1:
        cells = rows[0].findall(qn("w:tc"))
        widths = [COL0, COL1]
        for cell, w in zip(cells, widths):
            tcPr = cell.find(qn("w:tcPr"))
            if tcPr is None:
                continue
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.insert(0, tcW)
            tcW.set(qn("w:w"), str(w))
            tcW.set(qn("w:type"), "dxa")
            # Remove gridSpan entirely
            gs = tcPr.find(qn("w:gridSpan"))
            if gs is not None:
                tcPr.remove(gs)
    if len(rows) >= 2:
        cells = rows[1].findall(qn("w:tc"))
        if cells:
            cell = cells[0]
            tcPr = cell.find(qn("w:tcPr"))
            if tcPr is not None:
                tcW = tcPr.find(qn("w:tcW"))
                if tcW is None:
                    tcW = OxmlElement("w:tcW")
                    tcPr.insert(0, tcW)
                tcW.set(qn("w:w"), str(TOTAL))
                tcW.set(qn("w:type"), "dxa")
                gs = tcPr.find(qn("w:gridSpan"))
                if gs is None:
                    gs = OxmlElement("w:gridSpan")
                    tcPr.append(gs)
                gs.set(qn("w:val"), "2")

    # Set the table tblW to exact total
    if tblPr is not None:
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.insert(0, tblW)
        tblW.set(qn("w:w"), str(TOTAL))
        tblW.set(qn("w:type"), "dxa")
    return True


# ---------- footnote font fix ----------

def fix_footnote_fonts(docx_path: Path) -> dict:
    """Rewrite footnotes.xml: force Georgia, strip direct font sizes.

    Footnote runs that carry an explicit w:sz (e.g., 21 half-points = 10.5pt
    inherited from body) need that size stripped so the paragraph style's size
    (9pt from TME Footnote / footnote text) wins.
    """
    import zipfile, tempfile, os, re

    src = str(docx_path)
    with zipfile.ZipFile(src, "r") as z:
        if "word/footnotes.xml" not in z.namelist():
            return {"rfonts_rewritten": 0, "rfonts_injected": 0, "sz_stripped": 0}
        with z.open("word/footnotes.xml") as f:
            xml = f.read().decode("utf-8")

    # 1) Rewrite any existing w:rFonts to Georgia
    rewritten = 0
    def _rewrite_rfonts(m):
        nonlocal rewritten
        rewritten += 1
        return '<w:rFonts w:ascii="Georgia" w:hAnsi="Georgia" w:cs="Georgia"/>'
    xml = re.sub(r"<w:rFonts\b[^/]*/>", _rewrite_rfonts, xml)

    # 2) Strip direct w:sz / w:szCs so footnote runs inherit 9pt from paragraph style
    stripped = 0
    def _count_sub(pattern, s):
        nonlocal stripped
        s2, n = re.subn(pattern, "", s)
        stripped += n
        return s2
    xml = _count_sub(r"<w:sz\b[^/]*/>", xml)
    xml = _count_sub(r"<w:szCs\b[^/]*/>", xml)

    # 3) Inject Georgia rFonts into rPr blocks that lack one
    injected = 0
    def _inject_rfonts(m):
        nonlocal injected
        inner = m.group(1)
        if "<w:rFonts" in inner:
            return m.group(0)
        injected += 1
        return f'<w:rPr>{inner}<w:rFonts w:ascii="Georgia" w:hAnsi="Georgia" w:cs="Georgia"/></w:rPr>'
    xml = re.sub(r"<w:rPr>(.*?)</w:rPr>", _inject_rfonts, xml, flags=re.DOTALL)

    # Write back into the zip
    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(tmp_fd)
    with zipfile.ZipFile(src, "r") as zin, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/footnotes.xml":
                data = xml.encode("utf-8")
            zout.writestr(item, data)
    shutil.move(tmp_path, src)
    return {"rfonts_rewritten": rewritten, "rfonts_injected": injected, "sz_stripped": stripped}


# ---------- reference run cleanup ----------

def strip_reference_run_formatting(doc) -> int:
    """On every TME Reference paragraph, strip direct run-level font size,
    bold, and custom rFonts so the paragraph style (Georgia 10.5pt) wins."""
    n = 0
    for p in doc.paragraphs:
        if p.style.name != "TME Reference":
            continue
        for r in p.runs:
            rPr = r._r.find(qn("w:rPr"))
            if rPr is None:
                continue
            for tag in ("w:sz", "w:szCs", "w:b", "w:bCs", "w:rFonts"):
                for el in rPr.findall(qn(tag)):
                    rPr.remove(el)
                    n += 1
    return n


# ---------- main ----------

def main() -> None:
    # Only create the backup the first time so re-runs don't clobber it with
    # the already-fixed version.
    if not BACKUP.exists():
        shutil.copy2(STARTER, BACKUP)
        print(f"Backed up to {BACKUP.name}")
    else:
        print(f"Backup already at {BACKUP.name} (preserving)")

    doc = Document(str(STARTER))

    update_styles(doc)
    print("Updated style definitions (body, captions, footnote, block quote, list, headings)")

    bq_count = remap_block_quotes(doc)
    print(f"Remapped {bq_count} paragraphs -> TME Block Quote")

    caption_stats = fix_caption_classifications(doc)
    print(f"Reclassified captions: {caption_stats}")

    list_n = clear_list_direct_spacing(doc)
    print(f"Cleared direct spacing on {list_n} List Paragraph instances")

    t_count = fix_content_tables(doc)
    print(f"Centered {t_count} content tables, set cantSplit + tblHeader")

    masthead_ok = fix_masthead_grid(doc)
    print(f"Masthead grid rewrite: {'ok' if masthead_ok else 'skipped'}")

    ref_stripped = strip_reference_run_formatting(doc)
    print(f"Stripped {ref_stripped} run-level format attrs on TME Reference paragraphs")

    doc.save(str(STARTER))

    # footnote font fix mutates the zip directly; do it after save.
    fn_stats = fix_footnote_fonts(STARTER)
    print(f"Footnote rewrite: {fn_stats}")

    print("\nDone. Open the starter in Word to verify.")


if __name__ == "__main__":
    main()
