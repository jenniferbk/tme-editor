"""Apply TME paragraph styles to the Moore starter after a 'Keep Source Formatting'
paste. Also deletes the duplicated cover-content block that came from pasting
Moore_resolved.docx's own cover material.

Run after pasting the Moore body into TME_Moore_2026_starter.docx.
"""
import re
import shutil
from pathlib import Path

from docx import Document


STARTER = Path('/Users/jenniferkleiman/Documents/TME/TME_Moore_2026_starter.docx')
BACKUP = STARTER.with_suffix('.pre-styling.docx')

BODY_START_PREFIX = 'Conceptual analysis is a tool'
TITLE_PREFIX = 'Integration by Substitution:'

FIG_PAT = re.compile(r'^\s*Figure\s+\d', re.I)
TAB_PAT = re.compile(r'^\s*Table\s+\d', re.I)


def _first_nonempty_run(p):
    for r in p.runs:
        if r.text and r.text.strip():
            return r
    return None


def _style(doc, name: str):
    return doc.styles[name]


def main() -> None:
    shutil.copy2(STARTER, BACKUP)
    print(f'Backed up to {BACKUP}')

    doc = Document(STARTER)

    paras = list(doc.paragraphs)
    body_start_idx = next(
        (i for i, p in enumerate(paras) if p.text.startswith(BODY_START_PREFIX)),
        None,
    )
    if body_start_idx is None:
        raise RuntimeError(f'Could not find body anchor paragraph starting with {BODY_START_PREFIX!r}')

    dup_start_idx = None
    for i in range(body_start_idx - 1, -1, -1):
        p = paras[i]
        t = p.text.strip()
        if t.startswith(TITLE_PREFIX) and p.style.name != 'TME Title':
            dup_start_idx = i
            break

    deleted = 0
    if dup_start_idx is not None:
        for p in paras[dup_start_idx:body_start_idx]:
            p._element.getparent().remove(p._element)
            deleted += 1
        print(f'Deleted {deleted} duplicated cover paragraphs (idx {dup_start_idx}..{body_start_idx - 1})')

    paras = list(doc.paragraphs)
    body_start_idx = next(
        i for i, p in enumerate(paras) if p.text.startswith(BODY_START_PREFIX)
    )

    stats = {
        'TME Body': 0, 'TME H1': 0, 'TME Figure Caption': 0,
        'TME Table Caption': 0, 'TME Reference': 0, 'skipped_empty': 0,
    }

    for p in paras[body_start_idx:]:
        t = p.text.strip()
        if not t:
            stats['skipped_empty'] += 1
            continue
        src_style = p.style.name if p.style else ''

        if src_style in ('TMEReference', 'TME Reference'):
            p.style = _style(doc, 'TME Reference')
            stats['TME Reference'] += 1
            continue
        if src_style == 'Caption':
            if TAB_PAT.match(t):
                p.style = _style(doc, 'TME Table Caption')
                stats['TME Table Caption'] += 1
            else:
                p.style = _style(doc, 'TME Figure Caption')
                stats['TME Figure Caption'] += 1
            continue
        if src_style == 'EndNoteBibliographyTitle':
            p.style = _style(doc, 'TME H1')
            stats['TME H1'] += 1
            continue
        if src_style == 'ListParagraph':
            p.style = _style(doc, 'TME Body')
            stats['TME Body'] += 1
            continue
        if src_style in ('Normal', '', 'Default Paragraph Font'):
            first = _first_nonempty_run(p)
            is_bold = first is not None and first.bold
            is_short = len(t) < 150
            if is_bold and is_short:
                p.style = _style(doc, 'TME H1')
                stats['TME H1'] += 1
            else:
                p.style = _style(doc, 'TME Body')
                stats['TME Body'] += 1

    doc.save(STARTER)
    print('\nStyle application complete:')
    for k, v in stats.items():
        print(f'  {k:22s} {v}')


if __name__ == '__main__':
    main()
