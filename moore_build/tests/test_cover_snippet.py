from pathlib import Path

import pytest
from docx import Document

from moore_pipeline.cover_snippet import build_moore_cover_snippet


def test_cover_snippet_is_generated(tmp_path):
    # Use the already-prepared headshots from assets dir (regenerate if needed)
    from moore_pipeline.headshots import prepare_all_headshots
    TME = Path("/Users/jenniferkleiman/Documents/TME")
    assets = tmp_path / "assets"
    headshots = prepare_all_headshots(
        moore_src=TME / "Moore_Kevin.avif",
        yasuda_src=TME / "Yasuda_Sohei.jpg",
        wong_src=TME / "Wong_Webster.jpg",
        out_dir=assets,
    )

    out = tmp_path / "cover.docx"
    build_moore_cover_snippet(out_path=out, headshots=headshots)
    assert out.exists()

    doc = Document(str(out))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for cell in t._cells:
            for p in cell.paragraphs:
                all_text += "\n" + p.text
    # Also check footer tables (cover footer is in the section footer slot)
    for s in doc.sections:
        for t in s.footer.tables:
            for cell in t._cells:
                for p in cell.paragraphs:
                    all_text += "\n" + p.text

    # Core content present:
    assert "RESEARCH ARTICLE" in all_text
    assert "Vol. 34, No. 1 (2026)" in all_text
    assert "Integration by Substitution" in all_text
    assert "Kevin C. Moore" in all_text
    assert "Sohei Yasuda" in all_text
    assert "Webster Wong" in all_text
    assert "University of Georgia" in all_text
    assert "kvcmoore@uga.edu" in all_text
    assert "Conceptual Analysis" in all_text
    assert "Integration by Substitution" in all_text
    assert "CC BY 4.0" in all_text
