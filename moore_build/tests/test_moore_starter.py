from pathlib import Path

import pytest
from docx import Document

from moore_pipeline.moore_starter import build_moore_starter


def test_moore_starter_is_generated(tmp_path):
    from moore_pipeline.headshots import prepare_all_headshots
    TME = Path("/Users/jenniferkleiman/Documents/TME")
    assets = tmp_path / "assets"
    headshots = prepare_all_headshots(
        moore_src=TME / "Moore_Kevin.avif",
        yasuda_src=TME / "Yasuda_Sohei.jpg",
        wong_src=TME / "Wong_Webster.jpg",
        out_dir=assets,
    )

    out = tmp_path / "TME_Moore_2026_starter.docx"
    build_moore_starter(out_path=out, headshots=headshots)
    assert out.exists()


def test_moore_starter_has_cover_content(tmp_path):
    from moore_pipeline.headshots import prepare_all_headshots
    TME = Path("/Users/jenniferkleiman/Documents/TME")
    assets = tmp_path / "assets"
    headshots = prepare_all_headshots(
        moore_src=TME / "Moore_Kevin.avif",
        yasuda_src=TME / "Yasuda_Sohei.jpg",
        wong_src=TME / "Wong_Webster.jpg",
        out_dir=assets,
    )

    out = tmp_path / "TME_Moore_2026_starter.docx"
    build_moore_starter(out_path=out, headshots=headshots)

    doc = Document(str(out))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for cell in t._cells:
            for p in cell.paragraphs:
                all_text += "\n" + p.text
    # Also check footer tables (cover footer is in the section footer slot)
    for s in doc.sections:
        for t in s.footer.tables:
            for cell in t._cells:
                for p in cell.paragraphs:
                    all_text += "\n" + p.text

    # Core cover content present:
    assert "RESEARCH ARTICLE" in all_text
    assert "Vol. 34, No. 1 (2026)" in all_text
    assert "Integration by Substitution" in all_text
    assert "Kevin C. Moore" in all_text
    assert "Sohei Yasuda" in all_text
    assert "Webster Wong" in all_text
    assert "University of Georgia" in all_text
    assert "kvcmoore@uga.edu" in all_text
    assert "Conceptual Analysis" in all_text
    assert "CC BY 4.0" in all_text


def test_moore_starter_has_three_sections(tmp_path):
    """Document must have 3 sections: masthead/tagline (0), cover body (1), article body (2)."""
    from moore_pipeline.headshots import prepare_all_headshots
    TME = Path("/Users/jenniferkleiman/Documents/TME")
    assets = tmp_path / "assets"
    headshots = prepare_all_headshots(
        moore_src=TME / "Moore_Kevin.avif",
        yasuda_src=TME / "Yasuda_Sohei.jpg",
        wong_src=TME / "Wong_Webster.jpg",
        out_dir=assets,
    )

    out = tmp_path / "TME_Moore_2026_starter.docx"
    build_moore_starter(out_path=out, headshots=headshots)

    doc = Document(str(out))
    assert len(doc.sections) == 3, (
        f"Expected 3 sections (masthead, cover body, article body), got {len(doc.sections)}"
    )
    # Section 0: full-bleed (zero top/left/right margins) for masthead + tagline;
    # bottom_margin is 0.5" to give the cover footer room to render.
    s0 = doc.sections[0]
    assert s0.left_margin.inches == 0.0
    assert s0.right_margin.inches == 0.0
    assert s0.top_margin.inches == 0.0
    assert abs(s0.bottom_margin.inches - 0.5) < 0.01
    # Section 1: normal margins for cover body content
    s1 = doc.sections[1]
    assert abs(s1.left_margin.inches - 0.5) < 0.01
    assert abs(s1.right_margin.inches - 0.5) < 0.01
    assert abs(s1.top_margin.inches - 0.3) < 0.01
    assert abs(s1.bottom_margin.inches - 0.3) < 0.01


def test_moore_starter_body_section_headers(tmp_path):
    """Body section (section index 1) must have running headers set correctly."""
    from moore_pipeline.headshots import prepare_all_headshots
    TME = Path("/Users/jenniferkleiman/Documents/TME")
    assets = tmp_path / "assets"
    headshots = prepare_all_headshots(
        moore_src=TME / "Moore_Kevin.avif",
        yasuda_src=TME / "Yasuda_Sohei.jpg",
        wong_src=TME / "Wong_Webster.jpg",
        out_dir=assets,
    )

    out = tmp_path / "TME_Moore_2026_starter.docx"
    build_moore_starter(out_path=out, headshots=headshots)

    doc = Document(str(out))
    # Body section is sections[2] (0=masthead, 1=cover body, 2=article body)
    body_section = doc.sections[2]

    recto_text = "\n".join(p.text for p in body_section.header.paragraphs)
    verso_text = "\n".join(p.text for p in body_section.even_page_header.paragraphs)

    assert "Integration by Substitution" in recto_text, (
        f"Recto header missing short title. Got: {repr(recto_text)}"
    )
    assert "Moore, Yasuda, & Wong" in verso_text, (
        f"Verso header missing author cite. Got: {repr(verso_text)}"
    )


def test_moore_starter_has_placeholder_paragraph(tmp_path):
    """Body section must contain the placeholder paragraph styled as TME Body."""
    from moore_pipeline.headshots import prepare_all_headshots
    TME = Path("/Users/jenniferkleiman/Documents/TME")
    assets = tmp_path / "assets"
    headshots = prepare_all_headshots(
        moore_src=TME / "Moore_Kevin.avif",
        yasuda_src=TME / "Yasuda_Sohei.jpg",
        wong_src=TME / "Wong_Webster.jpg",
        out_dir=assets,
    )

    out = tmp_path / "TME_Moore_2026_starter.docx"
    build_moore_starter(out_path=out, headshots=headshots)

    doc = Document(str(out))
    all_para_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Paste Moore article body here" in all_para_text
