"""End-to-end: run the pipeline entry point and verify all three artifacts exist
and have plausible properties."""
import subprocess
import sys
from pathlib import Path

import pytest
from docx import Document


TME = Path("/Users/jenniferkleiman/Documents/TME")
MOORE_BUILD = TME / "moore_build"


@pytest.fixture(scope="module", autouse=True)
def _run_pipeline():
    # Use sys.executable so the venv python is used (not the system python3)
    subprocess.run(
        [sys.executable, "src/build_moore_cover.py"],
        cwd=str(MOORE_BUILD),
        check=True,
    )
    yield





def test_resolved_docx_exists():
    p = MOORE_BUILD / "intermediate" / "Moore_resolved.docx"
    assert p.exists()


def test_resolved_docx_has_no_endnote_markers():
    p = MOORE_BUILD / "intermediate" / "Moore_resolved.docx"
    doc = Document(str(p))
    all_text = "\n".join(para.text for para in doc.paragraphs)
    assert "ADDIN EN.CITE" not in all_text


def test_three_framed_headshots_exist():
    for name in ("moore", "yasuda", "wong"):
        p = MOORE_BUILD / "assets" / f"{name}_framed.jpg"
        assert p.exists()


def test_starter_docx_exists():
    p = TME / "TME_Moore_2026_starter.docx"
    assert p.exists()


def test_starter_contains_moore_content():
    p = TME / "TME_Moore_2026_starter.docx"
    doc = Document(str(p))
    all_text = "\n".join(para.text for para in doc.paragraphs)
    for table in doc.tables:
        for cell in table._cells:
            for para in cell.paragraphs:
                all_text += "\n" + para.text
    # Also check footer tables (cover footer is in the section footer slot)
    for s in doc.sections:
        for table in s.footer.tables:
            for cell in table._cells:
                for para in cell.paragraphs:
                    all_text += "\n" + para.text
    assert "Integration by Substitution" in all_text
    assert "Kevin C. Moore" in all_text
    assert "kvcmoore@uga.edu" in all_text
    assert "Dec 3, 2024" in all_text  # received date
    assert "CC BY 4.0" in all_text


def test_starter_has_two_sections():
    p = TME / "TME_Moore_2026_starter.docx"
    doc = Document(str(p))
    assert len(doc.sections) >= 3, (
        f"Expected 3 sections (masthead, cover body, article body), got {len(doc.sections)}"
    )


def test_starter_body_section_has_running_headers():
    p = TME / "TME_Moore_2026_starter.docx"
    doc = Document(str(p))
    # Body section is sections[2]: 0=masthead (zero margins), 1=cover body, 2=article body
    body_section = doc.sections[2]
    recto_text = "\n".join(para.text for para in body_section.header.paragraphs)
    verso_text = "\n".join(para.text for para in body_section.even_page_header.paragraphs)
    assert "Integration by Substitution" in recto_text
    assert "Moore, Yasuda, & Wong" in verso_text
